"""How many LLM calls each request type is allowed to make, and with what budget.

Latency work is only safe if the call count is pinned by tests: a hard grounding
failure must degrade to verified evidence instead of paying for a second generation,
and a general question must never touch the legal retrieval stack.
"""

import io

import pytest
from docx import Document as WordDocument

from app.config import get_settings
from app.services import ai_agent
from app.services.grounding import validate_cited_answer, validate_legal_answer
from app.services.llm import llm

LEGAL_SOURCES = [{
    "citation_number": 1,
    "document_id": "n1",
    "document_name": "Raqobat to‘g‘risida (O‘RQ-850)",
    "article_or_clause": "19-modda. Raqobatga qarshi kelishuvlar",
    "display_label": "19-modda",
    "url": "https://lex.uz/docs/6518381",
    "excerpt": "19-modda. Raqobatga qarshi kelishuvlarni tuzish taqiqlanadi.",
    "full_excerpt": "19-modda. Raqobatga qarshi kelishuvlarni tuzish taqiqlanadi.",
    "evidence_type": "nhh",
}]

DOCUMENT_SOURCES = [{
    "citation_number": 1,
    "document_id": "d1",
    "document_name": "hisobot.pdf",
    "article_or_clause": "2. Ko‘rsatkichlar",
    "url": "/api/documents/d1/download",
    "excerpt": "Jami murojaatlar: 128 ta. Jarayonda: 24 ta.",
    "evidence_type": "document",
}]


class CallRecorder:
    """Counts every generation the request path performs."""

    def __init__(self, monkeypatch, *, text=None, structured=None):
        self.calls: list[dict] = []
        self._text = text if isinstance(text, list) else [text]
        self._structured = structured if isinstance(structured, list) else [structured]

        async def fake_generate(system, user, **options):
            self.calls.append({"kind": "text", "options": options, "user": user})
            value = self._text.pop(0) if len(self._text) > 1 else self._text[0]
            if isinstance(value, Exception):
                raise value
            return value

        async def fake_structured(system, user, schema, **options):
            self.calls.append({"kind": "structured", "options": options, "user": user})
            value = self._structured.pop(0) if len(self._structured) > 1 else self._structured[0]
            if isinstance(value, Exception):
                raise value
            return value

        monkeypatch.setattr(llm, "generate", fake_generate)
        monkeypatch.setattr(llm, "generate_structured", fake_structured)

    @property
    def count(self) -> int:
        return len(self.calls)

    def budget(self, index: int = 0) -> int | None:
        return self.calls[index]["options"].get("max_tokens")


def blocks(text: str, source_ids=("L1",)) -> dict:
    return {"answer_blocks": [{"text": text, "source_ids": list(source_ids)}]}


def law_docx() -> bytes:
    stream = io.BytesIO()
    document = WordDocument()
    for line in ("19-modda. Raqobatga qarshi kelishuvlarni taqiqlash",
                 "Raqobatga qarshi kelishuvlarni tuzish taqiqlanadi."):
        document.add_paragraph(line)
    document.save(stream)
    return stream.getvalue()


# --- general chat --------------------------------------------------------------

def test_general_question_makes_one_call_and_skips_the_legal_stack(
        client, xodim_headers, monkeypatch):
    recorder = CallRecorder(monkeypatch, text="Umumiy javob tayyor.")

    def forbidden(*_args, **_kwargs):
        raise AssertionError("Umumiy savol huquqiy qidiruvni ishga tushirmasligi kerak")

    monkeypatch.setattr(ai_agent, "search_async", forbidden)
    monkeypatch.setattr(ai_agent, "filter_legal_topic", forbidden)
    monkeypatch.setattr(ai_agent, "distinct_source_chunks", forbidden)
    monkeypatch.setattr(ai_agent, "expand_article_sources", forbidden)

    response = client.post("/api/chat", headers=xodim_headers,
                           json={"question": "Bugungi ish rejamni uch bandda tuz.",
                                 "legal": False})
    assert response.status_code == 200, response.text
    body = response.json()
    assert body["effective_mode"] == "general"
    assert body["sources"] == []
    assert recorder.count == 1
    assert recorder.calls[0]["kind"] == "text"


def test_general_chat_uses_the_small_completion_budget(client, xodim_headers, monkeypatch):
    recorder = CallRecorder(monkeypatch, text="Qisqa javob.")
    client.post("/api/chat", headers=xodim_headers,
                json={"question": "Bugungi uchrashuv uchun agenda tuz.", "legal": False})
    assert recorder.budget() == get_settings().llm_max_tokens_general == 512


def test_legal_intent_still_auto_routes_in_auto_mode(client, admin_headers,
                                                    xodim_headers, monkeypatch):
    created = client.post(
        "/api/nhh", headers=admin_headers,
        data={"title": "Raqobat to‘g‘risida", "category": "Qonun",
              "source_url": "https://lex.uz/auto-route"},
        files={"file": ("qonun.docx", law_docx())},
    )
    assert created.status_code == 201
    recorder = CallRecorder(monkeypatch, structured=blocks(
        "Raqobatga qarshi kelishuvlarni tuzish taqiqlanadi"))
    response = client.post("/api/chat", headers=xodim_headers, json={
        "question": "Raqobatga qarshi kelishuvlar qaysi moddada tartibga solingan?",
        "mode": "auto",
    })
    assert response.status_code == 200
    body = response.json()
    assert body["routed_to_legal"] is True
    assert body["effective_mode"] == "legal"
    assert body["sources"], "huquqiy niyat aniqlansa manba ko‘rsatilishi kerak"
    # Routing costs at most one generation, and deterministic legal-fact extraction
    # may answer this class of question without any call at all.
    assert recorder.count <= 1


# --- legal generation ----------------------------------------------------------

@pytest.mark.anyio
async def test_successful_legal_answer_makes_exactly_one_call(monkeypatch):
    recorder = CallRecorder(monkeypatch, structured=blocks(
        "Raqobatga qarshi kelishuvlarni tuzish taqiqlanadi"))
    result = await ai_agent.generate_grounded_legal(
        "system", "prompt", LEGAL_SOURCES, question="kelishuvlar taqiqi")
    assert result.result_kind == "ok"
    assert recorder.count == 1
    assert recorder.budget() == get_settings().llm_max_tokens_legal


@pytest.mark.anyio
@pytest.mark.parametrize("bad_answer,code", [
    # invented article number
    ("Ushbu masala 77-modda bilan tartibga solinadi", "article_unsupported"),
    # unsupported numeric fact
    ("Kelishuvlar uchun 87 foizlik chegara belgilangan", "number_unsupported"),
    # invented law identifier
    ("O‘RQ-999 hujjatida shunday belgilangan", "document_id_unsupported"),
    # fabricated quotation attributed to the source
    ('Manbada “ushbu norma butunlay bekor qilingan deb hisoblanadi” deyilgan', "quote_unsupported"),
])
async def test_hard_grounding_failure_skips_the_correction_call(monkeypatch, bad_answer, code):
    """A factual failure must go straight to verified evidence, not a second call."""
    recorder = CallRecorder(monkeypatch, structured=blocks(bad_answer))
    result = await ai_agent.generate_grounded_legal(
        "system", "prompt", LEGAL_SOURCES, question="kelishuvlar taqiqi")
    assert result.result_kind == "source_matches"
    assert recorder.count == 1, "qattiq grounding xatosi ikkinchi chaqiriqqa arzimaydi"
    assert result.failure_reason == "validation_failed"
    # The verified extractive answer is served instead.
    assert "19-modda" in result.answer and "[1]" in result.answer
    assert code in validate_legal_answer(bad_answer + " [1]", LEGAL_SOURCES).codes


@pytest.mark.anyio
async def test_repairable_schema_failure_still_earns_one_correction_call(monkeypatch):
    """Evidence is fine and only the envelope was wrong, so retrying is worth it."""
    recorder = CallRecorder(monkeypatch, structured=[
        {"answer_blocks": []},  # schema-shaped failure
        blocks("Raqobatga qarshi kelishuvlarni tuzish taqiqlanadi"),
    ])
    result = await ai_agent.generate_grounded_legal(
        "system", "prompt", LEGAL_SOURCES, question="kelishuvlar taqiqi")
    assert result.result_kind == "ok"
    assert recorder.count == 2


@pytest.mark.anyio
async def test_drafting_keeps_the_large_budget(monkeypatch):
    recorder = CallRecorder(monkeypatch, structured=blocks(
        "Raqobatga qarshi kelishuvlarni tuzish taqiqlanadi"))
    await ai_agent.generate_grounded_legal(
        "system", "prompt", LEGAL_SOURCES, question="loyiha", budget_kind="drafting")
    assert recorder.budget() == get_settings().llm_max_tokens_drafting == 3200


# --- document analysis ---------------------------------------------------------

@pytest.mark.anyio
async def test_document_analysis_budget_and_single_call(monkeypatch):
    recorder = CallRecorder(monkeypatch, text="Jami murojaatlar 128 ta. [1]")
    result = await ai_agent.generate_grounded_document("system", "prompt", DOCUMENT_SOURCES)
    assert result.result_kind == "ok"
    assert recorder.count == 1
    assert recorder.budget() == get_settings().llm_max_tokens_document


@pytest.mark.anyio
async def test_document_hard_grounding_failure_skips_the_correction_call(monkeypatch):
    recorder = CallRecorder(monkeypatch, text="Jami murojaatlar 999 ta. [1]")
    result = await ai_agent.generate_grounded_document("system", "prompt", DOCUMENT_SOURCES)
    assert result.result_kind == "source_matches"
    assert recorder.count == 1
    assert "number_unsupported" in validate_cited_answer(
        "Jami murojaatlar 999 ta. [1]", DOCUMENT_SOURCES).codes


# --- deterministic paths -------------------------------------------------------

def test_deterministic_document_calculation_makes_no_llm_call(client, xodim_headers,
                                                              monkeypatch):
    from reportlab.pdfgen.canvas import Canvas

    stream = io.BytesIO()
    canvas = Canvas(stream)
    canvas.drawString(72, 780, "Jami murojaatlar: 128 ta")
    canvas.drawString(72, 760, "Jarayonda: 24 ta")
    canvas.save()
    uploaded = client.post("/api/documents", headers=xodim_headers,
                           data={"category": "oddiy"},
                           files={"file": ("hisobot.pdf", stream.getvalue())})
    assert uploaded.status_code == 201, uploaded.text

    recorder = CallRecorder(monkeypatch, text="LLM ishlatilmasligi kerak")
    response = client.post(f"/api/documents/{uploaded.json()['id']}/analyze",
                           headers=xodim_headers,
                           json={"operation": "qa", "question": "necha foizi jarayonda?"})
    assert response.status_code == 200, response.text
    assert "18,75" in response.json()["answer"]
    assert recorder.count == 0, "deterministik hisob LLMsiz bajarilishi kerak"


def test_deterministic_legal_threshold_makes_no_llm_call(client, admin_headers,
                                                         xodim_headers, monkeypatch):
    stream = io.BytesIO()
    document = WordDocument()
    document.add_paragraph("13-modda. Ustun mavqe")
    document.add_paragraph(
        "Xo‘jalik yurituvchi subyektning tovar bozoridagi ulushi qirq foizni va undan "
        "ortiqni tashkil etsa, uning holati ustun mavqe deb e’tirof etiladi."
    )
    document.save(stream)
    created = client.post(
        "/api/nhh", headers=admin_headers,
        data={"title": "Raqobat to‘g‘risida", "category": "Qonun",
              "source_url": "https://lex.uz/threshold"},
        files={"file": ("qonun.docx", stream.getvalue())},
    )
    assert created.status_code == 201

    recorder = CallRecorder(monkeypatch, structured=blocks("ishlatilmasligi kerak"))
    response = client.post("/api/chat", headers=xodim_headers, json={
        "question": "Bozor ulushi 35 foiz bo‘lgan korxona avtomatik ustun mavqe hisoblanadimi?",
        "legal": True,
    })
    assert response.status_code == 200, response.text
    assert response.json()["result_kind"] == "ok"
    assert recorder.count == 0, "chegara taqqoslash deterministik bo‘lishi kerak"
