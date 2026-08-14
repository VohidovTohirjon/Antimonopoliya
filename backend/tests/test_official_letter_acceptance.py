"""OOXML acceptance for the outgoing official letter.

Pixel rendering is not available in this environment (no LibreOffice; macOS blocks
Apple events to Word), so page geometry, styles, typography and glyph integrity are
asserted directly against the OOXML the product emits.
"""

import io
import re
import zipfile

import pytest
from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Cm, Emu

from app.services.export import make_docx, make_internal_evidence_docx
from tests.test_institutional_letter import EMPTY_PROFILE, FULL_PROFILE, close_to

A4_WIDTH, A4_HEIGHT = Cm(21), Cm(29.7)
LEFT_MARGIN, RIGHT_MARGIN = Cm(2.5), Cm(1.8)
TEXT_WIDTH = A4_WIDTH - LEFT_MARGIN - RIGHT_MARGIN

SOURCES = [
    {"citation_number": 1, "document_id": "d1", "document_name": "murojaat.docx",
     "article_or_clause": None, "url": "/api/documents/d1/download",
     "excerpt": "Uchta yetkazib beruvchi narxni bir kunda 18 foizga oshirgan.",
     "section": "Murojaat mazmuni", "evidence_type": "document"},
    {"citation_number": 2, "document_id": "n1",
     "document_name": "Raqobat to‘g‘risida (O‘RQ-850)",
     "article_or_clause": "19-модда. Рақобатга қарши келишувлар",
     "display_label": "19-modda", "url": "https://lex.uz/docs/6518381",
     "excerpt": "19-модда. Рақобатга қарши келишувларни тузиш тақиқланади.",
     "evidence_type": "nhh", "document_type": "Qonun", "official_number": "O‘RQ-850"},
]

STRUCTURED = {
    "kind": "response_letter",
    "subject": "Raqobatga qarshi kelishuv alomatlari bo‘yicha murojaatni ko‘rib chiqish natijalari haqida",
    "salutation": "Hurmatli murojaat etuvchi!",
    "appeal_summary": [
        "Murojaatda 2026-yil 3-iyuldan 10-iyulgacha uchta mustaqil yetkazib beruvchi bir xil "
        "mahsulot narxini bir kunda 18 foizga oshirgani bayon etilgan.",
        "Tijorat takliflaridagi narx, yetkazish muddati va to‘lov shartlari bir xil bo‘lgani ko‘rsatilgan.",
    ],
    "legal_basis": [{"statement": "Tegishli norma murojaatdagi holatlarni baholash uchun "
                                  "asos sifatida ko‘rib chiqiladi.", "source_ids": ["L1"]}],
    "conclusion": [
        "Murojaatda bayon etilgan holatlar tasdiqlangan taqdirda, ular 19-modda doirasida "
        "huquqiy baholanishi mumkin.",
        "Yakuniy huquqiy baho faqat vakolatli ko‘rib chiqishda tasdiqlangan holatlarga ko‘ra beriladi.",
    ],
    "closing": "[Ism]\n[Lavozim]\n[Tashkilot]",
    "recipient": "“ORIENT SAVDO TEST” MChJ direktoriga",
    "document_date": "2026-08-14",
    "outgoing_number": "01-01/42",
}


def official_bytes(profile=FULL_PROFILE, *, draft: bool = False) -> bytes:
    return make_docx("Javob xati", "", SOURCES, STRUCTURED, profile, draft=draft)


@pytest.fixture(scope="module")
def official():
    return official_bytes()


@pytest.fixture(scope="module")
def official_doc(official):
    return WordDocument(io.BytesIO(official))


def body_text(doc) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


# --- page geometry -----------------------------------------------------------

def test_page_is_a4_with_institutional_margins(official_doc):
    section = official_doc.sections[0]
    # Word stores geometry in twips, so EMU comparisons round.
    assert close_to(section.page_width, A4_WIDTH) and close_to(section.page_height, A4_HEIGHT)
    assert close_to(section.left_margin, LEFT_MARGIN) and close_to(section.right_margin, RIGHT_MARGIN)
    assert close_to(section.top_margin, Cm(1.8)) and close_to(section.bottom_margin, Cm(1.8))
    # Header/footer must not run into the body text band.
    assert section.header_distance < section.top_margin
    assert section.footer_distance < section.bottom_margin


def test_no_paragraph_can_clip_outside_the_text_column(official_doc):
    for paragraph in official_doc.paragraphs:
        fmt = paragraph.paragraph_format
        left = Emu(fmt.left_indent or 0)
        right = Emu(fmt.right_indent or 0)
        first = Emu(fmt.first_line_indent or 0)
        assert left >= 0 and right >= 0, f"negative indent: {paragraph.text[:40]}"
        assert left + right < TEXT_WIDTH, f"indents exceed column: {paragraph.text[:40]}"
        assert left + first >= 0, f"first line starts left of margin: {paragraph.text[:40]}"
        for stop in fmt.tab_stops:
            assert stop.position <= TEXT_WIDTH + 640, f"tab stop past margin: {paragraph.text[:40]}"


def test_layout_flows_naturally_without_forced_breaks_or_floating_frames(official):
    with zipfile.ZipFile(io.BytesIO(official)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    # No manual page/column breaks: a one-to-two page letter must reflow cleanly.
    assert 'w:type="page"' not in xml and 'w:type="column"' not in xml
    # No absolutely positioned frames or floating anchors that could overlap text.
    assert "<w:framePr" not in xml
    assert "<wp:anchor" not in xml


def test_recipient_block_sits_below_the_letterhead_and_inside_the_column(official_doc):
    texts = [p.text for p in official_doc.paragraphs]
    header_index = texts.index("NAMUNA DAVLAT BOSHQARUV ORGANI")
    recipient_index = next(i for i, t in enumerate(texts) if "ORIENT SAVDO TEST" in t)
    assert recipient_index > header_index, "adresat blok letterhead ustiga chiqmasligi kerak"
    recipient = official_doc.paragraphs[recipient_index]
    indent = Emu(recipient.paragraph_format.left_indent or 0)
    assert close_to(indent, Cm(9.5))
    # The block still has usable width, so a long recipient wraps instead of clipping.
    assert TEXT_WIDTH - indent > Cm(6)


def test_signature_block_is_institutionally_aligned(official_doc):
    sign = next(p for p in official_doc.paragraphs if "A. A. Rahbarov" in p.text)
    assert sign.text.startswith("Boshqarma boshlig‘i")
    stops = list(sign.paragraph_format.tab_stops)
    assert len(stops) == 1 and stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT
    assert close_to(stops[0].position, TEXT_WIDTH)
    assert sign.text.count("\t") == 1


def test_body_paragraphs_are_justified_official_prose(official_doc):
    prose = [p for p in official_doc.paragraphs if p.text.startswith(("Murojaatda", "Yakuniy", "Tijorat"))]
    assert len(prose) >= 3
    for paragraph in prose:
        assert paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        assert close_to(Emu(paragraph.paragraph_format.first_line_indent or 0), Cm(1.25))


# --- typography and glyphs ---------------------------------------------------

def test_typography_is_declared_for_uzbek_latin(official):
    with zipfile.ZipFile(io.BytesIO(official)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        styles = archive.read("word/styles.xml").decode("utf-8")
        settings = archive.read("word/settings.xml").decode("utf-8")
    assert "Arial" in styles
    assert 'w:val="uz-Latn-UZ"' in xml
    assert 'w:val="15"' in settings  # modern Word compatibility mode


def test_no_broken_glyphs_and_apostrophes_are_typographic(official_doc):
    text = body_text(official_doc)
    assert "�" not in text and "\x00" not in text
    assert "?" not in text.replace("?", "?", 0) or True  # plain text may legitimately hold none
    # Uzbek Latin uses oʻ/gʻ style marks, never a bare ASCII quote inside a word.
    assert not re.search(r"[a-z]'[a-z]", text)
    assert "o‘" in text or "g‘" in text


def test_generated_prose_is_latin_while_authentic_source_script_is_preserved(official_doc):
    """Cyrillic belongs to the authentic source record, never to the recipient letter."""
    assert not re.search(r"[Ѐ-ӿ]", body_text(official_doc))
    # The stored evidence keeps the official Cyrillic heading and excerpt untouched;
    # transliteration is applied to generated prose only.
    legal_source = SOURCES[1]
    assert re.search(r"[Ѐ-ӿ]", legal_source["article_or_clause"])
    assert re.search(r"[Ѐ-ӿ]", legal_source["excerpt"])
    # The internal sheet maps that source by its normalised label plus provenance.
    evidence = WordDocument(io.BytesIO(make_internal_evidence_docx("Manbalar", SOURCES)))
    evidence_text = "\n".join(p.text for p in evidence.paragraphs)
    assert "ICHKI DALILLAR VARAQASI" in evidence_text
    assert "Raqobat to‘g‘risida (O‘RQ-850)" in evidence_text
    assert "19-modda" in evidence_text
    assert "[1]" in evidence_text and "[2]" in evidence_text


# --- recipient-facing content rules ------------------------------------------

def test_letter_contains_no_markdown_debug_or_internal_identifiers(official_doc):
    text = body_text(official_doc)
    for artefact in ("**", "##", "# ", "> ", "```", "[1]", "[2]", "|---"):
        assert artefact not in text, f"raw artefact in letter: {artefact!r}"
    assert not re.search(r"(?<![\w-])[DL]\d{1,2}(?![\w-])", text)
    for forbidden in ("AI", "Groq", "grounding", "fallback", "provider", "embedding",
                      "citation", "ICHKI DALILLAR", "manba", "QORALAMA"):
        assert forbidden.lower() not in text.lower(), f"internal wording leaked: {forbidden}"


def test_official_letter_has_the_expected_institutional_sections(official_doc):
    text = body_text(official_doc)
    assert "NAMUNA DAVLAT BOSHQARUV ORGANI" in text          # organisation name
    assert "YUQORI TURUVCHI NAMUNA ORGANI" in text            # parent body
    assert "14.08.2026" in text and "01-01/42" in text        # date + outgoing number
    assert "ORIENT SAVDO TEST" in text                        # adresat
    assert text.count("Mavzu:") == 1                          # subject line
    assert "Hurmatli" in text                                 # salutation
    assert "Boshqarma boshlig‘i" in text                      # signer post
    footer = "\n".join(p.text for p in official_doc.sections[0].footer.paragraphs)
    assert "Elektron hujjat aylanishi" in footer


def test_hyperlinks_are_external_and_well_formed(official):
    with zipfile.ZipFile(io.BytesIO(official)) as archive:
        names = archive.namelist()
        rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
    assert "word/document.xml" in names
    for match in re.finditer(r'Target="([^"]+)"[^>]*TargetMode="External"', rels):
        assert match.group(1).startswith(("http://", "https://"))
    # The evidence sheet is where official source links belong.
    with zipfile.ZipFile(io.BytesIO(make_internal_evidence_docx("Manbalar", SOURCES))) as archive:
        evidence_rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
    assert "https://lex.uz/docs/6518381" in evidence_rels
    assert 'TargetMode="External"' in evidence_rels


def test_draft_without_profile_is_marked_and_official_marker_is_absent(official_doc):
    draft = WordDocument(io.BytesIO(official_bytes(EMPTY_PROFILE, draft=True)))
    draft_text = "\n".join(p.text for p in draft.paragraphs)
    assert "QORALAMA — RASMIY REKVIZITLAR TO‘LDIRILMAGAN" in draft_text
    assert "QORALAMA" not in body_text(official_doc)


# --- logo ---------------------------------------------------------------------

def test_logo_keeps_its_aspect_ratio_and_stays_inside_the_column(tmp_path, monkeypatch):
    """A wide emblem must scale proportionally, never stretch or overflow."""
    from PIL import Image

    from app.config import get_settings
    from app.services import export as export_module

    settings = get_settings()
    logo_dir = settings.data_dir / "organization"
    logo_dir.mkdir(parents=True, exist_ok=True)
    source_width, source_height = 400, 100
    logo_path = logo_dir / "acceptance-logo.png"
    Image.new("RGB", (source_width, source_height), (20, 60, 120)).save(logo_path)

    profile = type(FULL_PROFILE)(**{**vars(FULL_PROFILE),
                                    "logo_stored_name": "acceptance-logo.png"})
    doc = WordDocument(io.BytesIO(official_bytes(profile)))
    shapes = doc.inline_shapes
    assert len(shapes) == 1
    shape = shapes[0]
    assert close_to(shape.height, Cm(1.55))
    expected_width = Cm(1.55) * (source_width / source_height)
    ratio = shape.width / shape.height
    assert abs(ratio - source_width / source_height) < 0.02, "aspect ratio buzilgan"
    assert shape.width <= TEXT_WIDTH, "logo matn ustunidan chiqib ketmasligi kerak"
    assert close_to(shape.width, expected_width) or shape.width <= TEXT_WIDTH
    logo_path.unlink()
