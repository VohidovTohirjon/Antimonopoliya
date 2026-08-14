import io
from decimal import Decimal

from docx import Document as WordDocument
from app.services.legal_facts import deterministic_legal_fact_answer, parse_legal_source
from app.services.legal_intent import legal_concepts
from app.services.rag import filter_legal_topic


ARTICLE_13 = (
    "13-modda. Ustun mavqe "
    "Tovar yoki moliya bozorida xo‘jalik yurituvchi subyektning yoxud shaxslar guruhining "
    "raqobatlashuvchi xo‘jalik yurituvchi subyektlarga bog‘liq bo‘lmagan holda o‘z faoliyatini "
    "amalga oshirish va raqobatning holatiga hal qiluvchi ta’sir ko‘rsatish imkonini beradigan "
    "holati ustun mavqedir. "
    "Tovar yoki moliya bozorida quyidagilar ustun mavqe deb e’tirof etiladi: "
    "xo‘jalik yurituvchi subyektning yoki shaxslar guruhining raqobatchilari mavjud bo‘lmasa; "
    "xo‘jalik yurituvchi subyektning yoki shaxslar guruhining bozordagi ulushi qirq foiz va undan ortiq bo‘lsa; "
    "xo‘jalik yurituvchi subyekt tabiiy monopoliya subyekti bo‘lsa; "
    "xo‘jalik yurituvchi subyektga mutlaq yoki eksklyuziv huquq berilgan bo‘lsa. "
    "Oxirgi kalendar yil yakunlariga ko‘ra tovarlarni realizatsiya qilishdan olingan tushumi "
    "bazaviy hisoblash miqdorining o‘ttiz ming baravaridan kam bo‘lgan xo‘jalik yurituvchi "
    "subyektning yoki shaxslar guruhining holati ustun mavqe deb e’tirof etilmaydi "
    "(bundan tabiiy monopoliya subyektlari, mahsulot narxlari davlat tomonidan tartibga "
    "solinadigan xo‘jalik yurituvchi subyektlar mustasno). "
    "Ustun mavqeni e’tirof etish tartibi O‘zbekiston Respublikasi Vazirlar Mahkamasi tomonidan belgilanadi."
)

OTHER_ARTICLES = (
    "4-modda. Asosiy tushunchalar Ushbu Qonunda asosiy tushunchalar qo‘llaniladi.",
    "14-modda. Ustun muzokara kuchi Ustun muzokara kuchi bitim tarafining boshqa tarafga "
    "bog‘liq bo‘lmagan holda bitim shartlariga bir tomonlama ta’sir ko‘rsatish imkoniyatidir.",
    "18-modda. Ustun mavqeni suiiste’mol qilishni taqiqlash Ustun mavqeni egallab turgan "
    "xo‘jalik yurituvchi subyektning raqobatni cheklaydigan va boshqa shaxslarning manfaatlarini "
    "kamsitadigan harakatlari taqiqlanadi.",
    "19-modda. Raqobatga qarshi kelishuvlar Raqobatni cheklashga olib keladigan yoki olib "
    "kelishi mumkin bo‘lgan kelishuvlar taqiqlanadi.",
)

MANDATORY_QUERIES = (
    ("Raqobat to‘g‘risidagi qonunda ustun mavqeni aniqlash mezonlarini top", {"13-modda"}),
    ("Qaysi hollarda xo‘jalik yurituvchi subyekt bozorda ustun mavqega ega deb e’tirof etiladi?", {"13-modda"}),
    ("Korxonaning bozordagi dominant holatini aniqlash uchun qanday mezonlar qo‘llaniladi?", {"13-modda"}),
    ("Bozordagi ulush necha foizdan boshlab ustun mavqe mezoniga tushadi va boshqa shartlar qanday?", {"13-modda"}),
    ("Raqobatchisi mavjud bo‘lmagan korxona avtomatik ravishda ustun mavqega egami? Qonundagi shartlarni ko‘rsating.", {"13-modda"}),
    ("Bozor ulushi 35 foiz bo‘lgan korxona avtomatik ustun mavqe hisoblanadimi?", {"13-modda"}),
    ("Bozor ulushi 40 foiz bo‘lsa-chi?", {"13-modda"}),
    ("Ustun mavqe va ustun muzokara kuchi bir xil tushunchami?", {"13-modda", "14-modda"}),
    ("Ustun mavqeni suiiste’mol qilish nima?", {"18-modda"}),
    ("Dominant korxona raqobatga qarshi kelishuv qilsa qaysi normalar tegishli?", {"13-modda", "19-modda"}),
)


def _source_13() -> dict:
    return {
        "citation_number": 1,
        "article_or_clause": "13-modda. Ustun mavqe",
        "excerpt": ARTICLE_13,
        "full_excerpt": ARTICLE_13,
        "document_name": "Raqobat to‘g‘risidagi qonun",
    }


def _nhh_docx() -> bytes:
    stream = io.BytesIO()
    document = WordDocument()
    for paragraph in (OTHER_ARTICLES[0], ARTICLE_13, *OTHER_ARTICLES[1:]):
        document.add_paragraph(paragraph)
    document.save(stream)
    return stream.getvalue()


def _upload_law(client, admin_headers):
    response = client.post(
        "/api/nhh", headers=admin_headers,
        data={
            "title": "Raqobat to‘g‘risidagi qonun",
            "category": "Qonun",
            "source_url": "https://lex.uz/docs/official-reliability-test",
        },
        files={"file": ("raqobat-qonuni.docx", _nhh_docx())},
    )
    assert response.status_code == 201, response.text


def _articles(result: dict) -> set[str]:
    labels = set()
    for source in result["sources"]:
        value = source.get("article_or_clause") or source.get("display_label") or ""
        for number in (13, 14, 18, 19):
            if str(number) in value:
                labels.add(f"{number}-modda")
    return labels


def test_generic_legal_fact_model_reads_retrieved_source_and_threshold():
    facts = parse_legal_source(_source_13())
    assert facts is not None
    assert Decimal("40") in facts.percentages
    assert len(facts.clauses) >= 4

    for value in ("0", "20", "35", "39.9"):
        answer, sources = deterministic_legal_fact_answer(
            f"Korxona ulushi {value}% bo‘lsa dominantmi?", [_source_13()],
        )
        assert f"{value.replace('.', ',')}%" in answer
        assert "40%" in answer and "qanoatlantirmaydi" in answer
        assert sources
    for value in ("40", "45", "100"):
        answer, sources = deterministic_legal_fact_answer(
            f"Korxona ulushi {value}% bo‘lsa dominantmi?", [_source_13()],
        )
        assert f"{value}%" in answer
        assert "40%" in answer and "qanoatlantiradi" in answer
        assert sources


def test_dominance_synonyms_route_to_article_13_not_definitions():
    class Candidate:
        def __init__(self, article_clause, text):
            self.article_clause = article_clause
            self.text = text

    article4 = Candidate(OTHER_ARTICLES[0].split(" ", 2)[0] + " Asosiy tushunchalar", OTHER_ARTICLES[0])
    article13 = Candidate("13-modda. Ustun mavqe", ARTICLE_13)
    for question in (
        "dominant holat mezonlari", "dominant mavqe qachon yuzaga keladi",
        "dominant korxona qanday aniqlanadi", "bozorda hukmron holat shartlari",
        "bozordagi ustunlik mezonlari", "ustun holat qaysi hollarda e’tirof etiladi",
    ):
        concepts = legal_concepts(question)
        assert concepts.dominant, question
        assert filter_legal_topic([article4, article13], question) == [article13]


def test_previously_unseen_law_is_indexed_retrieved_and_cited(client, admin_headers, xodim_headers):
    stream = io.BytesIO()
    document = WordDocument()
    document.add_paragraph(
        "77-modda. Yashil infratuzilma reyestri. Reyestrga kiritish uchun hududning kamida "
        "ellik besh foizi ko‘kalamzorlashtirilgan bo‘lishi; suv tejovchi tizim ishlashi; "
        "yillik ekologik hisobot e’lon qilinishi kerak."
    )
    document.save(stream)
    created = client.post(
        "/api/nhh", headers=admin_headers,
        data={"title": "Sinov ekologik reyestr qoidalari", "category": "Nizom",
              "source_url": "https://lex.uz/docs/unseen-synthetic-uat"},
        files={"file": ("unseen-law.docx", stream.getvalue())},
    )
    assert created.status_code == 201, created.text
    response = client.post(
        "/api/chat", headers=xodim_headers,
        json={"question": "Yashil infratuzilma reyestriga kirish uchun qanday shartlar bor?",
              "legal": True},
    )
    assert response.status_code == 200, response.text
    result = response.json()
    assert result["sources"]
    assert any("77" in (source.get("article_or_clause") or "") for source in result["sources"])
    assert "[1]" in result["answer"]
