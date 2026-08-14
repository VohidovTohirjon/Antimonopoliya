"""Regression cover for the institutional outgoing-letter format and local seed policy.

These lock in the behaviour that separates a real official letter from an internal
report: localized correspondence dates, a signature block that reaches the right
margin, stationery ink instead of the in-app brand colour, and a seed command that
can never reset an existing account's password.
"""

import io
import re
import zipfile
from pathlib import Path
from types import SimpleNamespace

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Cm
from sqlalchemy import select

from app.config import get_settings
from app.database import SessionLocal
from app.models import Role, User
from app.security import hash_password, verify_password
from app.services.export import GREEN, _official_date, make_docx

PROFILE_FIELDS = (
    "organization_name", "organization_name_secondary", "short_name", "parent_organization",
    "address", "phone", "email", "website", "tax_id", "outgoing_prefix", "department",
    "letterhead_text", "footer_text", "signatory_name", "signatory_title",
    "qr_verification_url", "barcode_text",
)

FULL_PROFILE = SimpleNamespace(
    organization_name="Namuna davlat boshqaruv organi",
    organization_name_secondary="Sample State Authority",
    short_name="NDBO",
    parent_organization="Yuqori turuvchi namuna organi",
    address="100000, Toshkent shahri, Namuna ko‘chasi, 1-uy",
    phone="+998 71 000-00-00", email="devonxona@namuna.uz", website="namuna.uz",
    tax_id="300000000", outgoing_prefix="01-01", department="Huquqiy ta’minot boshqarmasi",
    letterhead_text="Rasmiy yozishmalar bo‘limi",
    footer_text="Elektron hujjat aylanishi tizimi orqali yuborilgan.",
    signatory_name="A. A. Rahbarov", signatory_title="Boshqarma boshlig‘i",
    qr_verification_url="", barcode_text="", logo_stored_name=None,
)

EMPTY_PROFILE = SimpleNamespace(
    **{name: "" for name in PROFILE_FIELDS}, logo_stored_name=None
)


def letter(profile, *, draft: bool, document_date: str = "2026-08-13") -> bytes:
    structured = {
        "kind": "response_letter",
        "subject": "Murojaatni ko‘rib chiqish natijalari haqida",
        "salutation": "Hurmatli murojaat etuvchi!",
        "appeal_summary": ["Murojaatda bayon etilgan holatlar ko‘rib chiqildi."],
        "legal_basis": [],
        "conclusion": ["Yakuniy baho vakolatli ko‘rib chiqishda beriladi."],
        "closing": "[Ism]\n[Lavozim]\n[Tashkilot]",
        "recipient": "“Namuna savdo” MChJ direktoriga",
        "document_date": document_date,
        "outgoing_number": "01-01/42",
    }
    return make_docx("Javob xati", "", [], structured, profile, draft=draft)


def paragraph_named(doc: WordDocument, needle: str):
    return next(p for p in doc.paragraphs if needle in p.text)


def close_to(measured, expected) -> bool:
    """Word stores positions in twips, so an EMU round-trip is lossy by design."""
    return abs(int(measured) - int(expected)) <= 640  # < 0.002 cm


def test_correspondence_date_is_localized_and_missing_date_stays_a_placeholder():
    assert _official_date("2026-08-13") == "13.08.2026"
    assert _official_date("") == "[sana]"
    assert _official_date(None) == "[sana]"
    # A value the product did not produce as ISO is passed through untouched
    # rather than silently reinterpreted into another calendar order.
    assert _official_date("13.08.2026") == "13.08.2026"

    doc = WordDocument(io.BytesIO(letter(FULL_PROFILE, draft=False)))
    reference = paragraph_named(doc, "№")
    assert "13.08.2026" in reference.text
    assert "2026-08-13" not in reference.text


def test_signature_block_reaches_the_right_margin_instead_of_a_default_tab():
    doc = WordDocument(io.BytesIO(letter(FULL_PROFILE, draft=False)))
    sign = paragraph_named(doc, "A. A. Rahbarov")
    assert sign.text.startswith("Boshqarma boshlig‘i")
    stops = list(sign.paragraph_format.tab_stops)
    assert len(stops) == 1
    assert stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT
    # A4 width minus the official-letter side margins.
    assert close_to(stops[0].position, Cm(21) - Cm(2.5) - Cm(1.8))


def test_recipient_is_an_indented_block_not_a_ragged_right_aligned_line():
    doc = WordDocument(io.BytesIO(letter(FULL_PROFILE, draft=False)))
    recipient = paragraph_named(doc, "Namuna savdo")
    assert recipient.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert close_to(recipient.paragraph_format.left_indent, Cm(9.5))


def test_official_letterhead_uses_stationery_ink_not_the_in_app_brand_colour():
    doc = WordDocument(io.BytesIO(letter(FULL_PROFILE, draft=False)))
    heading = paragraph_named(doc, "NAMUNA DAVLAT BOSHQARUV ORGANI")
    colours = {run.font.color.rgb for run in heading.runs if run.font.color.rgb}
    assert colours and GREEN not in colours

    with zipfile.ZipFile(io.BytesIO(letter(FULL_PROFILE, draft=False))) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    # The letterhead rule and organisation name must not carry the app's green.
    assert "12664E" not in xml


def test_draft_and_official_markers_stay_distinct_for_the_same_letter():
    draft_doc = WordDocument(io.BytesIO(letter(EMPTY_PROFILE, draft=True)))
    draft_text = "\n".join(p.text for p in draft_doc.paragraphs)
    assert "QORALAMA — RASMIY REKVIZITLAR TO‘LDIRILMAGAN" in draft_text

    reviewed = WordDocument(io.BytesIO(letter(FULL_PROFILE, draft=True)))
    reviewed_text = "\n".join(p.text for p in reviewed.paragraphs)
    assert "QORALAMA — YUBORISHDAN OLDIN VAKOLATLI XODIM TEKSHIRUVI" in reviewed_text

    official = WordDocument(io.BytesIO(letter(FULL_PROFILE, draft=False)))
    official_text = "\n".join(p.text for p in official.paragraphs)
    assert "QORALAMA" not in official_text


def test_internal_evidence_identifiers_never_reach_the_recipient_letter():
    """D1/L1 index the internal evidence sheet and are not recipient-facing."""
    structured = {
        "kind": "response_letter", "subject": "Sinov",
        "salutation": "Hurmatli murojaat etuvchi!",
        "appeal_summary": ["Murojaatda keltirilgan faktlar (D1) asosida holat ko‘rib chiqildi."],
        "legal_basis": [{"statement": "Norma [L1] doirasida baholanadi.", "source_ids": []}],
        "conclusion": ["Xulosa D1 va L2 dalillariga tayanadi. [2]"],
        "closing": "[Ism]", "recipient": "Qabul qiluvchi",
        "document_date": "2026-08-14", "outgoing_number": "01-01/42",
    }
    doc = WordDocument(io.BytesIO(
        make_docx("Javob xati", "", [], structured, FULL_PROFILE, draft=False)))
    body = "\n".join(p.text for p in doc.paragraphs)
    assert not re.search(r"(?<![\w-])[DL]\d{1,2}(?![\w-])", body)
    assert "[2]" not in body
    # Prose must survive the removal cleanly — no doubled spaces, no space
    # stranded before punctuation.
    assert "Murojaatda keltirilgan faktlar asosida holat ko‘rib chiqildi." in body
    prose = [p.text for p in doc.paragraphs
             if p.text.startswith(("Murojaatda", "Norma", "Xulosa"))]
    assert len(prose) == 3
    for line in prose:
        assert "  " not in line
        assert " ." not in line and " ," not in line


def test_no_migration_rewrites_stored_account_passwords():
    versions = Path(__file__).resolve().parents[1] / "alembic" / "versions"
    files = sorted(versions.glob("*.py"))
    assert files, "alembic revisions topilmadi"
    for path in files:
        body = path.read_text(encoding="utf-8")
        assert not re.search(r"password_hash\s*=", body), f"{path.name} parol hashini yozmoqda"
        assert "hash_password" not in body, f"{path.name} parol generatsiya qilmoqda"


def test_local_seed_password_is_configurable_and_never_overwrites_an_existing_account():
    from scripts.seed_demo_data import ensure_user

    assert len(get_settings().local_seed_password) >= 8

    with SessionLocal() as db:
        created = ensure_user(db, "seed_probe", "Seed sinovi", "SeedPass123!", Role.xodim)
        db.commit()
        assert verify_password("SeedPass123!", created.password_hash)

        original_hash = created.password_hash
        # A second seed run must return the same account untouched, so an
        # administrator-set password survives re-seeding.
        again = ensure_user(db, "seed_probe", "Seed sinovi", "AnotherPassword!", Role.xodim)
        db.commit()
        assert again.id == created.id
        assert again.password_hash == original_hash
        assert not verify_password("AnotherPassword!", again.password_hash)

        stored = db.scalar(select(User).where(User.username == "seed_probe"))
        assert stored.password_hash == original_hash
        assert stored.password_hash != hash_password("SeedPass123!")  # bcrypt salts per call
