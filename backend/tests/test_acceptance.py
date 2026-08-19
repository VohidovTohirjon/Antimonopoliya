import base64
import io
import re
import zipfile
from datetime import datetime, timedelta, timezone
from types import SimpleNamespace

from docx import Document as WordDocument
from fastapi import HTTPException
from openpyxl import Workbook
from reportlab.pdfgen.canvas import Canvas
from sqlalchemy import func, select

from app.database import SessionLocal
from app.models import AiHistory, AuditLog, Chunk, Document, OrganizationProfile, Task, User
from app.services.ai_agent import has_legal_intent
from app.services.grounding import (extractive_legal_fallback,
                                    repair_document_citations, used_sources,
                                    validate_cited_answer, validate_legal_answer)
from app.services.llm import llm
from app.services.export import make_docx
from app.services.rag import (_maintenance_penalty, _tokens, chunk_text,
                              filter_legal_topic, search)
from scripts.seed_demo_data import OPERATIONAL_FILES, SEED_KEY_PREFIX, TASKS, seed


def docx_bytes(text="Raqobat hujjati matni. Muhim qoida amal qiladi."):
    stream = io.BytesIO(); doc = WordDocument(); doc.add_paragraph(text)
    table = doc.add_table(rows=1, cols=2); table.cell(0, 0).text = "Band"; table.cell(0, 1).text = "Qiymat"
    doc.save(stream); return stream.getvalue()


def xlsx_bytes():
    stream = io.BytesIO(); wb = Workbook(); ws = wb.active; ws.title = "Hisobot"
    ws.append(["Ko‘rsatkich", "Qiymat"]); ws.append(["XLSX_UNIQUE_TEST_789", 12]); wb.save(stream); return stream.getvalue()


def pdf_bytes():
    stream = io.BytesIO(); canvas = Canvas(stream); canvas.drawString(72, 770, "PDF_UNIQUE_TEST_123"); canvas.save(); return stream.getvalue()


def upload(client, headers, filename, content, category="oddiy"):
    response = client.post("/api/documents", headers=headers, data={"category": category, "confidential": "true"}, files={"file": (filename, content)})
    assert response.status_code == 201, response.text
    return response.json()


def test_auth_valid_invalid_and_current_user(client):
    assert client.post("/api/auth/login", data={"username": "xodim", "password": "xato"}).status_code == 401
    response = client.post("/api/auth/login", data={"username": "xodim", "password": "Xodim123!"})
    assert response.status_code == 200
    assert response.json()["user"]["role"] == "xodim"
    headers = {"Authorization": "Bearer " + response.json()["access_token"]}
    assert client.get("/api/auth/me", headers=headers).json()["username"] == "xodim"
    assert client.post("/api/auth/logout", headers=headers).status_code == 204
    assert client.get("/api/auth/me", headers=headers).status_code == 401


def test_unauthenticated_and_role_protected_routes(client, xodim_headers):
    assert client.get("/api/documents").status_code == 401
    assert client.get("/api/history").status_code == 401
    assert client.post("/api/nhh", headers=xodim_headers).status_code == 403
    assert client.get("/api/roles", headers=xodim_headers).status_code == 403


def test_system_status_and_uzbek_script_normalization(client, xodim_headers):
    status = client.get("/api/ai/readiness", headers=xodim_headers)
    assert status.status_code == 200
    assert status.json()["legal_ready"] is False
    assert status.json()["status"] in {"ready", "preparing", "unavailable"}
    assert _tokens("Рақобат тўғрисидаги қонун") >= {"raqobat", "to'g'risidagi", "qonun"}
    assert "mavqe" in _tokens("ustun mavqeni aniqlash mezonlarini")
    chunks = chunk_text("13-модда. Устун мавқе\n\nТовар бозоридаги мезонлар.")
    assert chunks[0]["article"].startswith("13-модда")
    assert has_legal_intent("Raqobatni ta'minlash haqida O'zbekiston qonunchiligida qanday moddalar bor?")
    assert has_legal_intent("Қонуннинг 13-моддаси нимани белгилайди?")
    assert not has_legal_intent("Bugungi ishlarimni qanday rejalashtiraman?")
    assert not has_legal_intent("Bugungi ichki ish rejasini uch bandda tuzib ber.")
    for question in (
        "Ustun mavqedagi xo‘jalik yurituvchi subyekt uchun qanday cheklovlar mavjud?",
        "Raqobatga qarshi kelishuvlar qaysi moddada tartibga solingan?",
        "Savdolarda raqobatni cheklash bo‘yicha qonun nima deydi?",
    ):
        assert has_legal_intent(question)
    for question in (
        "Rahbar uchun ushbu matndan 5 bandli xulosa yoz.",
        "Bugungi uchrashuv uchun qisqa agenda tayyorla.",
        "Ushbu ikki fikrni taqqoslab ber.",
    ):
        assert not has_legal_intent(question)


def test_amendment_boilerplate_is_demoted_unless_requested():
    chunk = Chunk(article_clause="46-modda. Ayrim qonun hujjatlariga o‘zgartirish va qo‘shimchalar kiritish")
    assert _maintenance_penalty("Raqobatning asosiy qoidalari qaysilar?", chunk) > 0
    assert _maintenance_penalty("Qonunga kiritilgan o‘zgartirish va qo‘shimchalarni top", chunk) == 0


def test_named_legal_topic_drops_weakly_related_article():
    article_19 = SimpleNamespace(
        article_clause="19-модда. Рақобатга қарши келишувлар",
        text="Рақобатга қарши келишувларни тузиш тақиқланади.",
    )
    article_8 = SimpleNamespace(
        article_clause="8-модда. Норматив ҳужжатларнинг рақобатга таъсири",
        text="Рақобатни чекловчи нормаларнинг таъсири баҳоланади.",
    )
    result = filter_legal_topic(
        [article_19, article_8], "Raqobatni cheklovchi kelishuvlarga oid moddalarni top",
    )
    assert result == [article_19]


def test_rbac_admin_and_dashboard(client, xodim_headers, rahbar_headers):
    assert client.get("/api/users", headers=xodim_headers).status_code == 403
    assert client.get("/api/audit", headers=xodim_headers).status_code == 403
    assert client.get("/api/dashboard", headers=xodim_headers).status_code == 403
    dashboard = client.get("/api/dashboard", headers=rahbar_headers)
    assert dashboard.status_code == 200
    assert set(dashboard.json()) >= {"mavjud_muammolar", "kechikayotgan_topshiriqlar", "muhim_murojaatlar", "statistika", "xavfli_holatlar", "xodimlar_ish_yuklamasi"}


def test_upload_and_parse_all_supported_formats(client, xodim_headers):
    cases = [
        ("test.pdf", pdf_bytes(), "PDF_UNIQUE_TEST_123"),
        ("test.docx", docx_bytes("DOCX_UNIQUE_TEST_456 raqobat matni"), "DOCX_UNIQUE_TEST_456"),
        ("test.xlsx", xlsx_bytes(), "XLSX_UNIQUE_TEST_789 | 12"),
    ]
    for name, content, marker in cases:
        item = upload(client, xodim_headers, name, content)
        with SessionLocal() as db:
            stored = db.get(Document, item["id"])
            assert marker in stored.parsed_text
            assert db.scalar(select(Chunk).where(Chunk.document_id == stored.id))
    bad = client.post("/api/documents", headers=xodim_headers, files={"file": ("fake.pdf", b"not a pdf")})
    assert bad.status_code == 415
    mismatch = client.post("/api/documents", headers=xodim_headers, files={"file": ("wrong.pdf", docx_bytes())})
    assert mismatch.status_code == 415
    corrupt = io.BytesIO()
    with zipfile.ZipFile(corrupt, "w") as archive:
        archive.writestr("word/document.xml", "not xml")
    corrupt_response = client.post("/api/documents", headers=xodim_headers, files={"file": ("broken.docx", corrupt.getvalue())})
    assert corrupt_response.status_code == 422
    oversized = client.post("/api/documents", headers=xodim_headers, files={"file": ("huge.pdf", b"%PDF-" + b"0" * (20 * 1024 * 1024))})
    assert oversized.status_code == 413


def test_document_security_and_permission_filtered_rag(client, xodim_headers):
    item = upload(client, xodim_headers, "private.docx", docx_bytes("Maxfiy monopoliyaga oid matn"))
    other = client.post("/api/auth/login", data={"username": "other", "password": "Other123!"}).json()
    other_headers = {"Authorization": "Bearer " + other["access_token"]}
    assert client.get(f"/api/documents/{item['id']}", headers=other_headers).status_code == 403
    assert client.get(f"/api/documents/{item['id']}/download", headers=other_headers).status_code == 403
    assert client.post(f"/api/documents/{item['id']}/analyze", headers=other_headers, json={"operation": "summary"}).status_code == 403
    with SessionLocal() as db:
        user = db.scalar(select(User).where(User.username == "other"))
        assert search(db, "monopoliyaga", user, "document", item["id"]) == []

    public = client.post(
        "/api/documents", headers=xodim_headers, data={"category": "oddiy"},
        files={"file": ("public.docx", docx_bytes("Ochiq hujjat mazmuni"))},
    )
    assert public.status_code == 201
    public_id = public.json()["id"]
    listed = client.get("/api/documents", headers=other_headers).json()
    assert public_id in {value["id"] for value in listed}
    assert item["id"] not in {value["id"] for value in listed}
    assert client.get(f"/api/documents/{public_id}", headers=other_headers).status_code == 200


def test_document_analysis_all_operations(client, xodim_headers):
    item = upload(client, xodim_headers, "analysis.docx", docx_bytes())
    operations = [("summary", None), ("key_points", None), ("qa", "Muhim qoida nima?"), ("contradictions", None)]
    for operation, question in operations:
        response = client.post(f"/api/documents/{item['id']}/analyze", headers=xodim_headers, json={"operation": operation, "question": question})
        assert response.status_code == 200, response.text
        assert response.json()["answer"]
        if response.json()["sources"]:
            assert response.json()["sources"][0]["document_name"] == "analysis.docx"

    conflict = upload(
        client, xodim_headers, "conflict.docx",
        docx_bytes("Hisobot 15-avgust kuni topshiriladi. Boshqa bandda hisobot 20-avgust kuni topshiriladi."),
    )
    found = client.post(f"/api/documents/{conflict['id']}/analyze", headers=xodim_headers, json={"operation": "contradictions"})
    assert found.status_code == 200
    assert "15-avgust" in found.json()["answer"] and "20-avgust" in found.json()["answer"]


def test_document_analysis_rejects_orphan_citations(client, xodim_headers, monkeypatch):
    item = upload(client, xodim_headers, "grounded.docx", docx_bytes("Tasdiqlangan raqam 128 ta."))

    async def invented(_system: str, _user: str, **_options) -> str:
        return "Hujjatda 999 ta murojaat bor [99]."

    monkeypatch.setattr(llm, "generate", invented)
    response = client.post(
        f"/api/documents/{item['id']}/analyze", headers=xodim_headers,
        json={"operation": "summary"},
    )
    assert response.status_code == 200
    result = response.json()
    assert result["result_kind"] == "source_matches"
    assert "999" not in result["answer"] and "128" in result["answer"]
    assert result["warning"] and "tekshirilgan" in result["warning"]
    assert "citation" not in result["warning"].lower()
    citation_ids = [source["citation_number"] for source in result["sources"]]
    assert citation_ids == list(range(1, len(citation_ids) + 1))
    assert all(f"[{value}]" in result["answer"] for value in citation_ids)


def test_document_qa_calculates_explicit_percentage_and_refuses_missing_fact(
        client, xodim_headers):
    item = upload(
        client, xodim_headers, "ko‘rsatkichlar.docx",
        docx_bytes("Asosiy ko‘rsatkichlar. Jami murojaatlar: 128 ta. Jarayonda: 24 ta."),
    )
    calculated = client.post(
        f"/api/documents/{item['id']}/analyze", headers=xodim_headers,
        json={"operation": "qa", "question": "Asosiy ko‘rsatkichlarning necha foizi jarayonda?"},
    )
    assert calculated.status_code == 200, calculated.text
    assert "18,75%" in calculated.json()["answer"]
    assert "24 / 128 × 100" in calculated.json()["answer"]
    assert calculated.json()["structured"]["method"] == "percentage"
    total = client.post(
        f"/api/documents/{item['id']}/analyze", headers=xodim_headers,
        json={"operation": "qa", "question": "Jami murojaatlar soni nechta?"},
    )
    assert total.status_code == 200 and "128 ta" in total.json()["answer"]
    in_progress = client.post(
        f"/api/documents/{item['id']}/analyze", headers=xodim_headers,
        json={"operation": "qa", "question": "Jarayonda nechta murojaat bor?"},
    )
    assert in_progress.status_code == 200 and "24 ta" in in_progress.json()["answer"]
    monthly_item = upload(
        client, xodim_headers, "oylik.docx",
        docx_bytes("Oylik dinamika. Oy Jami Ko‘rib chiqilgan Jarayonda. Aprel 36 30 6 May 43 32 11 Iyun 49 34 15."),
    )
    monthly = client.post(
        f"/api/documents/{monthly_item['id']}/analyze", headers=xodim_headers,
        json={"operation": "qa", "question": "Iyun oyida nechta murojaat kelib tushgan?"},
    )
    assert monthly.status_code == 200 and "49 ta" in monthly.json()["answer"]
    assert monthly.json()["structured"]["method"] == "monthly_table"
    concise = client.post(
        f"/api/documents/{item['id']}/analyze", headers=xodim_headers,
        json={"operation": "qa", "question": "Hujjatni 1-2 gapda izohlang."},
    )
    assert concise.status_code == 200
    assert concise.json()["structured"]["method"] == "extractive_summary"
    assert 1 <= concise.json()["answer"].count("[") <= 2
    missing = client.post(
        f"/api/documents/{item['id']}/analyze", headers=xodim_headers,
        json={"operation": "qa", "question": "Tashkilot qachon tuzilgan?"},
    )
    assert missing.status_code == 200
    assert missing.json()["answer"] == "Bu ma’lumot tanlangan hujjatda topilmadi."
    assert missing.json()["sources"] == []


def test_document_citation_must_contain_claimed_numbers():
    sources = [
        {"citation_number": 1, "excerpt": "Aprel oyida jami 36 ta murojaat bo‘lgan."},
        {"citation_number": 2, "excerpt": "Iyun oyida jami 49 ta, jarayonda 15 ta murojaat bo‘lgan."},
    ]
    wrong = validate_cited_answer("Iyun oyida 49 ta, jarayonda 15 ta edi [1].", sources)
    assert wrong.valid is False
    assert any("49" in violation and "15" in violation for violation in wrong.violations)
    wrong_after_period = validate_cited_answer("Iyun oyida 49 ta, jarayonda 15 ta edi. [1]", sources)
    assert wrong_after_period.valid is False
    repaired = repair_document_citations("Iyun oyida 49 ta, jarayonda 15 ta edi. [1]", sources)
    assert repaired.endswith("[2]")
    assert validate_cited_answer(repaired, sources).valid is True
    right = validate_cited_answer("Iyun oyida 49 ta, jarayonda 15 ta edi [2].", sources)
    assert right.valid is True


def test_nhh_ingest_retrieve_metadata_reindex_and_no_fabrication(client, admin_headers, xodim_headers):
    response = client.post("/api/nhh", headers=admin_headers, data={"title": "Raqobat to‘g‘risidagi qonun", "category": "Qonun", "source_url": "https://lex.uz/test"}, files={"file": ("qonun.docx", docx_bytes("12-modda. Ustun mavqe mezonlari bozor ulushi asosida aniqlanadi."))})
    assert response.status_code == 201, response.text
    nhh = response.json(); assert nhh["indexed"] is True
    legal = client.post("/api/chat", headers=xodim_headers, json={"question": "Ustun mavqe mezonlari nima?", "legal": True})
    assert legal.status_code == 200
    source = legal.json()["sources"][0]
    assert source["citation_number"] == 1
    assert source["document_name"] == "Raqobat to‘g‘risidagi qonun"
    assert source["url"] == "https://lex.uz/test"
    assert "12-modda" in source["excerpt"]
    absent_article = client.post(
        "/api/chat", headers=xodim_headers,
        json={"question": "Raqobat to‘g‘risidagi qonunning 99-moddasi nima deydi?", "legal": True},
    )
    assert absent_article.status_code == 200
    assert absent_article.json()["result_kind"] == "no_sources"
    assert absent_article.json()["sources"] == []
    with SessionLocal() as db:
        before = db.scalar(select(func.count(Chunk.id)).where(Chunk.nhh_id == nhh["id"]))
    assert client.post(f"/api/nhh/{nhh['id']}/reindex", headers=admin_headers).status_code == 200
    with SessionLocal() as db:
        after = db.scalar(select(func.count(Chunk.id)).where(Chunk.nhh_id == nhh["id"]))
    assert before == after
    updated = client.patch(f"/api/nhh/{nhh['id']}", headers=admin_headers, json={"category": "Nizom", "source_url": "https://lex.uz/updated"})
    assert updated.status_code == 200 and updated.json()["category"] == "Nizom"
    invalid = client.patch(f"/api/nhh/{nhh['id']}", headers=admin_headers, json={"category": "Noma’lum"})
    assert invalid.status_code == 422
    assert client.delete(f"/api/nhh/{nhh['id']}", headers=admin_headers).status_code == 204


def test_every_supported_nhh_category_uses_the_same_validated_workflow(
        client, admin_headers, xodim_headers):
    categories = (
        "Qonun", "Prezident farmoni", "Prezident qarori",
        "Vazirlar Mahkamasi qarori", "Nizom", "Buyruq",
        "Idoraviy (ichki) hujjat",
    )
    created_ids: list[str] = []
    for position, category in enumerate(categories, start=1):
        data = {"title": f"{category} sinov hujjati", "category": category}
        if category != "Idoraviy (ichki) hujjat":
            data["source_url"] = f"https://lex.uz/category-test-{position}"
        response = client.post(
            "/api/nhh", headers=admin_headers, data=data,
            files={"file": (f"nhh-{position}.docx", docx_bytes(
                f"{position}-modda. Ushbu toifa yagona tekshirilgan indekslash jarayonidan o‘tadi."
            ))},
        )
        assert response.status_code == 201, response.text
        item = response.json()
        assert item["category"] == category and item["indexed"] is True
        created_ids.append(item["id"])
    for item_id in created_ids:
        assert client.delete(f"/api/nhh/{item_id}", headers=admin_headers).status_code == 204
    missing = client.post("/api/chat", headers=xodim_headers, json={"question": "Mavjud bo‘lmagan qonun?", "legal": True})
    assert missing.json()["sources"] == []
    assert missing.json()["result_kind"] == "corpus_empty"
    assert "bazasi hozircha bo‘sh" in missing.json()["answer"]


def test_auto_mode_routes_legal_intent_and_citations_stay_mapped(
        client, admin_headers, xodim_headers, monkeypatch):
    """Legal-intent inference now belongs to the opt-in "auto" mode.

    Explicit "general" is a contract and is covered by test_chat_mode_contract.
    """
    long_article = "12-modda. Ustun mavqe mezonlari. " + (
        "Tovar bozoridagi ulush va raqobat sharoitlari rasmiy mezon sifatida tekshiriladi. " * 35
    )
    created = client.post(
        "/api/nhh", headers=admin_headers,
        data={"title": "Raqobat to‘g‘risidagi qonun", "category": "Qonun",
              "source_url": "https://lex.uz/docs/official-test"},
        files={"file": ("qonun.docx", docx_bytes(long_article))},
    )
    assert created.status_code == 201, created.text

    async def citations_with_an_invalid_number(_system: str, _user: str, **_options) -> str:
        return "**12-modda** bo‘yicha rasmiy mezon topildi [1]. Uydirma manba [99]."

    monkeypatch.setattr(llm, "generate", citations_with_an_invalid_number)
    response = client.post(
        "/api/chat", headers=xodim_headers,
        json={"question": "O‘zbekiston qonunchiligida ustun mavqe qaysi modda bilan belgilanadi?",
              "mode": "auto"},
    )
    assert response.status_code == 200, response.text
    result = response.json()
    assert result["effective_mode"] == "legal"
    assert result["routed_to_legal"] is True
    assert result["result_kind"] == "ok"
    assert "[1]" in result["answer"] and "[99]" not in result["answer"]
    assert [source["citation_number"] for source in result["sources"]] == list(
        range(1, len(result["sources"]) + 1)
    )
    assert len(result["sources"]) == 1
    assert result["sources"][0]["url"] == "https://lex.uz/docs/official-test"


def test_general_work_question_stays_general(client, xodim_headers):
    response = client.post(
        "/api/chat", headers=xodim_headers,
        json={"question": "Bugungi ishlarimni qanday rejalashtiraman?", "legal": False},
    )
    assert response.status_code == 200
    assert response.json()["effective_mode"] == "general"
    assert response.json()["routed_to_legal"] is False
    assert response.json()["sources"] == []


def test_citation_validator_rejects_nonexistent_source_number():
    sources = [{"citation_number": 1, "document_name": "Raqobat to‘g‘risida (O‘RQ-850)",
                "article_or_clause": "19-modda", "excerpt": "19-modda. Kelishuvlar."}]
    result = validate_legal_answer("Da’vo [1]. Yana bir da’vo [99].", sources)
    assert result.valid is False
    assert any("citation" in value for value in result.violations)


def test_grounding_rejects_fabricated_legal_identifiers_and_normalizes_terminology():
    sources = [{
        "citation_number": 1,
        "document_name": "Raqobat to‘g‘risida (O‘RQ-850)",
        "article_or_clause": "19-модда. Рақобатга қарши келишувлар",
        "excerpt": "19-модда. Рақобатга қарши келишувларни тузиш тақиқланади.",
    }]
    fabricated = validate_legal_answer(
        "O‘zbekiston Respublikasi Raqobatni nazorat qilish to‘g‘risidagi qonun, "
        "2019-yil, №12-III, 12-modde talabiga ko‘ra taqiqlanadi [1].",
        sources,
    )
    assert fabricated.valid is False
    assert any("modda" in value for value in fabricated.violations)
    assert any("hujjat raqami" in value for value in fabricated.violations)
    assert any("hujjat nomi" in value for value in fabricated.violations)
    assert "12-modda" in fabricated.answer and "modde" not in fabricated.answer

    unicode_hyphen = validate_legal_answer("Uydirma 12‑modda talabi [1].", sources)
    assert unicode_hyphen.valid is False
    assert "12-modda" in unicode_hyphen.answer

    supported = validate_legal_answer("Kelishuvlar 19-modd bilan tartibga solingan [1].", sources)
    assert supported.valid is True
    assert "19-modda" in supported.answer
    assert re.search(r"19-modd(?:\s|$)", supported.answer) is None

    supported_unicode_id = validate_legal_answer("O‘RQ‑850 hujjatining 19‑modda talabi [1].", sources)
    assert supported_unicode_id.valid is True

    fabricated_draft_date = validate_legal_answer(
        "2099-yil 12-iyul. 19-modda talabi [1].", sources,
        additional_evidence="Murojaat 2026-yil 8-iyulda berilgan.",
    )
    assert fabricated_draft_date.valid is False
    assert any("2099" in value and "12" in value for value in fabricated_draft_date.violations)


def test_used_source_mapping_is_bidirectional_and_has_no_orphan_source():
    sources = [
        {"citation_number": 1, "document_name": "Qonun", "article_or_clause": "19-modda",
         "excerpt": "19-modda. Kelishuvlar."},
        {"citation_number": 2, "document_name": "Qonun", "article_or_clause": "8-modda",
         "excerpt": "8-modda. NHH ta’sirini baholash."},
    ]
    validation = validate_legal_answer("Kelishuvlar bo‘yicha asos [1].", sources)
    assert validation.valid is True
    displayed = used_sources(sources, validation.used_citation_ids)
    assert [source["citation_number"] for source in displayed] == [1]
    assert all(f"[{source['citation_number']}]" in validation.answer for source in displayed)


def test_extractive_fallback_cites_every_displayed_source():
    sources = [
        {"citation_number": 1, "document_name": "Qonun", "article_or_clause": "19-модда",
         "excerpt": "19-модда. Келишувлар."},
        {"citation_number": 2, "document_name": "Qonun", "article_or_clause": "29-модда",
         "excerpt": "29-модда. Савдолар."},
    ]
    answer, displayed = extractive_legal_fallback(sources, "AI izohi mavjud emas")
    # The invariant is displayed == cited. Evidence that survived topic filtering is
    # rendered rather than dropped, so a multi-article question keeps every article.
    assert displayed == sources
    assert "19-modda" in answer and "29-modda" in answer
    assert all(f"[{source['citation_number']}]" in answer for source in displayed)
    assert answer.count("[1]") == 1 and answer.count("[2]") == 1


def test_same_legal_query_has_stable_evidence_order_ten_times(client, admin_headers, xodim_headers):
    created = client.post(
        "/api/nhh", headers=admin_headers,
        data={"title": "Raqobat to‘g‘risida (O‘RQ-850)", "category": "Qonun",
              "source_url": "https://lex.uz/stable"},
        files={"file": ("qonun.docx", docx_bytes(
            "19-modda. Raqobatga qarshi kelishuvlar taqiqlanadi.\n\n"
            "29-modda. Savdolarda raqobatni cheklash taqiqlanadi.\n\n"
            "8-modda. Normativ hujjatlarning ta’siri baholanadi."
        ))},
    )
    assert created.status_code == 201
    evidence_sets = []
    for _ in range(10):
        result = client.post("/api/chat", headers=xodim_headers, json={
            "question": "Raqobatni cheklovchi kelishuvlarga oid moddalarni toping", "legal": True,
        }).json()
        evidence_sets.append(tuple((source["citation_number"], source["article_or_clause"])
                                   for source in result["sources"]))
    assert len(set(evidence_sets)) == 1


def test_operational_seed_is_idempotent_and_uses_real_document_model():
    args = SimpleNamespace(confirm_development=True, admin_username="admin", reset=False,
                           reset_only=False)
    seed(args)
    seed(args)
    with SessionLocal() as db:
        documents = list(db.scalars(select(Document).where(Document.filename.in_(OPERATIONAL_FILES))))
        tasks = list(db.scalars(select(Task).where(Task.seed_key.like(f"{SEED_KEY_PREFIX}%"))))
        assert len(documents) == len(OPERATIONAL_FILES)
        assert len(tasks) == len(TASKS)
        assert all(document.parsed_text for document in documents)
        assert all(db.scalar(select(Chunk).where(Chunk.document_id == document.id)) for document in documents)


def test_unrelated_legal_question_returns_no_sources_without_fabrication(client, admin_headers, xodim_headers):
    created = client.post(
        "/api/nhh", headers=admin_headers,
        data={"title": "Raqobat qoidalari", "category": "Qonun", "source_url": "https://lex.uz/raqobat"},
        files={"file": ("raqobat.docx", docx_bytes("12-modda. Ustun mavqe tovar bozori ulushi asosida belgilanadi."))},
    )
    assert created.status_code == 201
    answer = client.post("/api/chat", headers=xodim_headers, json={"question": "Mars sayyorasida kartoshka hosili qancha?", "legal": True})
    assert answer.status_code == 200
    assert answer.json()["result_kind"] == "no_sources"
    assert answer.json()["sources"] == []
    assert "yetarli huquqiy asos topilmadi" in answer.json()["answer"]


def test_grounded_draft_history_and_valid_docx_export(client, admin_headers, xodim_headers):
    missing_document = client.post("/api/drafts", headers=xodim_headers, json={"kind": "response_letter", "instruction": "Javob xati yoz"})
    assert missing_document.status_code == 422
    doc = upload(client, xodim_headers, "murojaat.docx", docx_bytes("Fuqaro raqobat shartlari haqida murojaat qildi."), "murojaat")
    client.post("/api/nhh", headers=admin_headers, data={"title": "Qonun", "category": "Qonun", "source_url": "https://lex.uz/qonun"}, files={"file": ("qonun.docx", docx_bytes("20-modda. Murojaat qonuniy tartibda ko‘rib chiqiladi."))})
    response = client.post("/api/drafts", headers=xodim_headers, json={"kind": "response_letter", "instruction": "Ushbu murojaatga javob xatini tayyorla", "document_id": doc["id"]})
    assert response.status_code == 200, response.text
    result = response.json(); assert result["sources"] and result["export_url"]
    export = client.get(result["export_url"], headers=xodim_headers)
    assert export.status_code == 200
    assert zipfile.is_zipfile(io.BytesIO(export.content))
    assert WordDocument(io.BytesIO(export.content)).paragraphs
    history = client.get("/api/history", headers=xodim_headers).json()
    assert any(item["operation"] == "response_letter" for item in history)

    note = client.post("/api/drafts", headers=xodim_headers, json={"kind": "info_note", "instruction": "Haftalik ma’lumotnoma tayyorla"})
    assert note.status_code == 200
    note_export = client.get(note.json()["export_url"], headers=xodim_headers)
    note_doc = WordDocument(io.BytesIO(note_export.content))
    assert note_export.status_code == 200 and note_doc.paragraphs[0].text == "Ma’lumotnoma"

    other = client.post("/api/auth/login", data={"username": "other", "password": "Other123!"}).json()
    other_headers = {"Authorization": "Bearer " + other["access_token"]}
    assert all(item["user_id"] == other["user"]["id"] for item in client.get("/api/history", headers=other_headers).json())
    assert client.get(result["export_url"], headers=other_headers).status_code == 403


def test_grounded_draft_survives_generation_rate_limit(
        client, admin_headers, xodim_headers, monkeypatch):
    document = upload(
        client, xodim_headers, "murojaat.docx",
        docx_bytes("Fuqaro ustun mavqe mezonlari haqida murojaat qildi."), "murojaat",
    )
    created = client.post(
        "/api/nhh", headers=admin_headers,
        data={"title": "Raqobat qonuni", "category": "Qonun",
              "source_url": "https://lex.uz/rate-limit-test"},
        files={"file": ("qonun.docx", docx_bytes(
            "12-modda. Ustun mavqe tovar bozoridagi ulush asosida tekshiriladi."
        ))},
    )
    assert created.status_code == 201

    async def rate_limited(_system: str, _user: str, **_options) -> str:
        raise HTTPException(503, "AI xizmati so‘rovlar chegarasiga yetdi")

    monkeypatch.setattr(llm, "generate", rate_limited)
    response = client.post("/api/drafts", headers=xodim_headers, json={
        "kind": "response_letter", "instruction": "Huquqiy javob xati tayyorla",
        "document_id": document["id"],
    })
    assert response.status_code == 200, response.text
    result = response.json()
    assert result["result_kind"] == "source_matches"
    assert "mas’ul xodim tahriri" in result["answer"]
    assert result["sources"][0]["citation_number"] == 1
    assert client.get(result["export_url"], headers=xodim_headers).status_code == 200


def test_tasks_overdue_authorization_and_dashboard(client, admin_headers, rahbar_headers, xodim_headers):
    users = client.get("/api/users", headers=admin_headers).json(); worker = next(u for u in users if u["username"] == "xodim")
    created = client.post("/api/tasks", headers=rahbar_headers, json={"title": "Muddatli topshiriq", "description": "Tekshiruv", "assigned_to": worker["id"], "deadline": (datetime.now(timezone.utc)-timedelta(days=1)).isoformat()})
    assert created.status_code == 201, created.text
    task = created.json(); assert task["is_overdue"] is True
    assert client.post("/api/tasks", headers=xodim_headers, json={"title": "Ruxsatsiz", "assigned_to": worker["id"], "deadline": datetime.now(timezone.utc).isoformat()}).status_code == 403
    visible = client.get("/api/tasks", headers=xodim_headers).json(); assert visible[0]["id"] == task["id"]
    assert client.patch(f"/api/tasks/{task['id']}", headers=xodim_headers, json={"status": "bajarildi"}).status_code == 200
    dashboard = client.get("/api/dashboard", headers=rahbar_headers).json()
    assert dashboard["statistika"]["bajarilgan_topshiriqlar"] == 1
    assert dashboard["kechikayotgan_topshiriqlar"] == []
    assert dashboard["xavfli_holatlar"] == []
    worker_load = next(item for item in dashboard["xodimlar_ish_yuklamasi"] if item["user_id"] == worker["id"])
    assert worker_load == {
        "user_id": worker["id"], "full_name": worker["full_name"], "jami": 1,
        "faol": 0, "bajarilgan": 1, "kechikkan": 0, "progress_percent": 100,
    }


def test_admin_user_management_and_audit(client, admin_headers, xodim_headers):
    create = client.post("/api/users", headers=admin_headers, json={"username": "newuser", "full_name": "Yangi xodim", "password": "Secret123!", "role": "xodim"})
    assert create.status_code == 201
    updated = client.patch(f"/api/users/{create.json()['id']}", headers=admin_headers, json={"role": "rahbar", "is_active": False})
    assert updated.json()["role"] == "rahbar" and updated.json()["is_active"] is False
    client.get("/api/documents", headers=xodim_headers)
    logs = client.get("/api/audit", headers=admin_headers)
    assert logs.status_code == 200
    assert any(item["path"] == "/api/documents" for item in logs.json())
    with SessionLocal() as db:
        assert db.scalar(select(AuditLog)) is not None


def test_admin_changes_revoke_sessions_and_last_admin_is_protected(client, admin_headers):
    me = client.get("/api/auth/me", headers=admin_headers).json()
    assert client.patch(f"/api/users/{me['id']}", headers=admin_headers, json={"role": "xodim"}).status_code == 422

    created = client.post("/api/users", headers=admin_headers, json={
        "username": "sessionuser", "full_name": "Sessiya xodimi", "password": "Secret123!", "role": "xodim",
    })
    token = client.post("/api/auth/login", data={"username": "sessionuser", "password": "Secret123!"}).json()["access_token"]
    user_headers = {"Authorization": "Bearer " + token}
    assert client.get("/api/auth/me", headers=user_headers).status_code == 200
    assert client.patch(f"/api/users/{created.json()['id']}", headers=admin_headers, json={"is_active": False}).status_code == 200
    assert client.get("/api/auth/me", headers=user_headers).status_code == 401


def test_admin_can_configure_and_remove_official_logo(client, admin_headers):
    logo = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
    )
    uploaded = client.post(
        "/api/organization-profile/logo", headers=admin_headers,
        files={"file": ("logo.png", logo, "image/png")},
    )
    assert uploaded.status_code == 200, uploaded.text
    assert uploaded.json()["logo_stored_name"].endswith(".png")
    with SessionLocal() as db:
        profile = db.scalar(select(OrganizationProfile))
        exported = make_docx("Javob xati", "", [], {
            "kind": "response_letter", "subject": "Sinov", "salutation": "Hurmatli murojaatchi,",
            "appeal_summary": ["Murojaat ko‘rib chiqildi."], "legal_basis": [],
            "conclusion": ["Loyiha vakolatli xodim tekshiruviga taqdim etiladi."],
        }, profile)
    rendered = WordDocument(io.BytesIO(exported))
    assert len(rendered.inline_shapes) == 1
    removed = client.delete("/api/organization-profile/logo", headers=admin_headers)
    assert removed.status_code == 200
    assert removed.json()["logo_stored_name"] is None


def test_confidential_document_text_never_reaches_external_provider(
        client, xodim_headers, monkeypatch):
    item = upload(
        client, xodim_headers, "maxfiy.docx",
        docx_bytes("MAXFIY_PROVIDERGA_CHIQMASIN_91827 ichki ma’lumot."),
    )
    calls: list[str] = []

    async def forbidden_provider(_system: str, user: str, **_options) -> str:
        calls.append(user)
        raise AssertionError("Maxfiy matn tashqi adapterga uzatildi")

    monkeypatch.setattr(llm, "generate", forbidden_provider)
    response = client.post(
        f"/api/documents/{item['id']}/analyze", headers=xodim_headers,
        json={"operation": "summary"},
    )
    assert response.status_code == 200, response.text
    assert response.json()["result_kind"] == "source_matches"
    assert calls == []


def test_official_export_requires_complete_profile_and_correspondence_fields(
        client, admin_headers, xodim_headers):
    document = upload(
        client, xodim_headers, "murojaat.docx",
        docx_bytes("Fuqaro raqobatga qarshi kelishuv haqida murojaat qildi."), "murojaat",
    )
    created = client.post(
        "/api/nhh", headers=admin_headers,
        data={"title": "Raqobat to‘g‘risida", "category": "Qonun",
              "official_number": "O‘RQ-850", "source_url": "https://lex.uz/official-export"},
        files={"file": ("qonun.docx", docx_bytes(
            "19-modda. Raqobatga qarshi kelishuvlarni tuzish taqiqlanadi."
        ))},
    )
    assert created.status_code == 201
    result = client.post("/api/drafts", headers=xodim_headers, json={
        "kind": "response_letter", "instruction": "Javob xatini tayyorla",
        "document_id": document["id"],
    }).json()
    assert result["official_export_url"] is None
    assert result["export_missing_fields"]
    forbidden = client.get(
        f"/api/history/{result['history_id']}/export?official=true", headers=xodim_headers,
    )
    assert forbidden.status_code == 422
    draft = WordDocument(io.BytesIO(client.get(result["export_url"], headers=xodim_headers).content))
    draft_text = "\n".join(paragraph.text for paragraph in draft.paragraphs)
    assert "QORALAMA" in draft_text and "[sana]" in draft_text

    profile = {
        "organization_name": "Sinov davlat tashkiloti",
        "organization_name_secondary": "Test State Organization",
        "short_name": "SDT", "parent_organization": "Yuqori turuvchi sinov organi",
        "address": "Toshkent shahri", "phone": "+998 71 000-00-00",
        "email": "devonxona@example.uz", "website": "example.uz", "tax_id": "123456789",
        "outgoing_prefix": "01-01", "department": "Huquqiy boshqarma",
        "letterhead_text": "Rasmiy yozishmalar bo‘limi", "footer_text": "Sinov footer matni",
        "signatory_name": "A. Rahbar", "signatory_title": "Tashkilot rahbari",
        "qr_verification_url": "", "barcode_text": "",
    }
    assert client.put("/api/organization-profile", headers=admin_headers, json=profile).status_code == 200
    completed = client.post("/api/drafts", headers=xodim_headers, json={
        "kind": "response_letter", "instruction": "Javob xatini tayyorla",
        "document_id": document["id"], "recipient": "Murojaat muallifiga",
        "document_date": "2026-08-13", "outgoing_number": "01-01/42",
    })
    assert completed.status_code == 200, completed.text
    completed_result = completed.json()
    assert completed_result["official_export_url"]
    official = client.get(completed_result["official_export_url"], headers=xodim_headers)
    assert official.status_code == 200
    official_doc = WordDocument(io.BytesIO(official.content))
    official_text = "\n".join(p.text for p in official_doc.paragraphs)
    assert "QORALAMA" not in official_text
    assert "[sana]" not in official_text and "[1]" not in official_text
    assert "O‘zbekiston Respublikasining “Raqobat to‘g‘risida”gi Qonunining 19-moddasi" in official_text
    assert "YUQORI TURUVCHI SINOV ORGANI" in official_text
    assert "TEST STATE ORGANIZATION" in official_text
    assert "STIR: 123456789" in official_text
    assert any("Sinov footer matni" in p.text for p in official_doc.sections[0].footer.paragraphs)
