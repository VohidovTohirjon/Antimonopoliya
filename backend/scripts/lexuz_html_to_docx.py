"""Convert a previously downloaded Lex.uz document page into an upload-ready DOCX."""

import argparse
import re
from pathlib import Path

from docx import Document
from lxml import html


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--source-url", required=True)
    args = parser.parse_args()

    page = html.fromstring(args.input.read_text(encoding="utf-8"))
    nodes = page.xpath(
        "//div[contains(concat(' ', normalize-space(@class), ' '), ' lx_elem ')]"
        "/div[@name and @id]"
    )
    paragraphs = [re.sub(r"\s+", " ", " ".join(node.itertext())).strip() for node in nodes]
    paragraphs = [value for value in paragraphs if value]
    if len(paragraphs) < 20 or sum(map(len, paragraphs)) < 10_000:
        raise SystemExit("Lex.uz sahifasidan yetarli rasmiy matn ajratilmadi")

    document = Document()
    document.core_properties.title = paragraphs[1] if len(paragraphs) > 1 else "Lex.uz NHH"
    document.add_heading(document.core_properties.title, 0)
    document.add_paragraph(f"Rasmiy manba: {args.source_url}")
    for value in paragraphs:
        if re.match(r"^\d+\s*-?\s*боб\.", value, re.I):
            document.add_heading(value, level=1)
        elif re.match(r"^\d+\s*-?\s*модда\.", value, re.I):
            document.add_heading(value, level=2)
        else:
            document.add_paragraph(value)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)


if __name__ == "__main__":
    main()
