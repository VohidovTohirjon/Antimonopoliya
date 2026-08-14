"""Explicit, idempotent operational sample-data seed.

This command never runs automatically. It generates real PDF/DOCX/XLSX files,
parses them with the production parser and indexes their parsed text.
"""

import argparse
import hashlib
import re
import io
import sys
from datetime import datetime, timedelta, timezone
from pathlib import Path

from docx import Document as WordDocument
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen.canvas import Canvas
from sqlalchemy import delete, select

BACKEND_ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = BACKEND_ROOT.parent
sys.path.insert(0, str(BACKEND_ROOT))

from app.config import get_settings
from app.database import SessionLocal
from app.models import Document, Role, Task, TaskStatus, User
from app.security import hash_password
from app.services.documents import MIMES, parse_file
from app.services.rag import index_document


OPERATIONAL_FILES = {
    "murojaat_raqobat.docx": ("docx", "murojaat"),
    "tahliliy_hisobot.pdf": ("pdf", "muhim"),
    "topshiriq_bayonnomasi.docx": ("docx", "oddiy"),
    "murojaatlar_statistikasi.xlsx": ("xlsx", "oddiy"),
}
SEED_KEY_PREFIX = "operational-sample:v2:"
TASKS = (
    ("Murojaat bo‘yicha huquqiy asoslarni tekshirish", TaskStatus.jarayonda, -2),
    ("Savdo bo‘yicha tahliliy ma’lumotnoma tayyorlash", TaskStatus.jarayonda, 3),
    ("NHH metama’lumotlarini tekshirish", TaskStatus.jarayonda, 5),
    ("Kechikayotgan murojaat holatini ko‘rib chiqish", TaskStatus.yangi, -1),
    ("Rahbar uchun haftalik qisqa hisobot tayyorlash", TaskStatus.yangi, 1),
    ("Kelishuvlar bo‘yicha dalillarni guruhlash", TaskStatus.yangi, 2),
    ("Murojaat ijrochisiga eslatma yuborish", TaskStatus.yangi, 7),
    ("Oylik statistik jadvalni tekshirish", TaskStatus.bajarildi, -4),
    ("Tahliliy hisobot ma’lumotlarini yig‘ish", TaskStatus.bajarildi, -3),
    ("Hujjat parsing natijasini tasdiqlash", TaskStatus.bajarildi, -1),
)


def make_appeal() -> bytes:
    stream = io.BytesIO()
    doc = WordDocument()
    doc.add_heading("Raqobat shartlari bo‘yicha murojaat", 1)
    doc.add_heading("Murojaat mazmuni", 2)
    doc.add_paragraph(
        "‘ORIENT SAVDO TEST’ MChJ vakili sifatida 2026-yil 3-iyuldan 10-iyulgacha uchta "
        "mustaqil yetkazib beruvchi bir xil mahsulot narxini bir kunda 18 foizga oshirganini ma’lum qilamiz."
    )
    doc.add_paragraph(
        "Yetkazib beruvchilarning tijorat takliflaridagi narx, yetkazish muddati va to‘lov shartlari "
        "bir xil bo‘lgan. 2026-yil 8-iyuldagi telefon suhbatida bir yetkazib beruvchi narxlar "
        "‘bozor ishtirokchilari bilan kelishilganini’ bildirgan."
    )
    doc.add_heading("So‘rov", 2)
    doc.add_paragraph(
        "Mazkur holatda raqobatga qarshi kelishuv alomatlari mavjudligini vakolat doirasida "
        "ko‘rib chiqish, zarur hujjatlar ro‘yxatini bildirish va huquqiy asoslangan javob berishni so‘raymiz."
    )
    doc.add_heading("Ilova sifatida mavjud dalillar", 2)
    for value in ("uchta tijorat taklifi", "narxlar taqqoslama jadvali", "telefon suhbati bo‘yicha ichki qayd"):
        doc.add_paragraph(value, style="List Bullet")
    doc.save(stream)
    return stream.getvalue()


def make_report() -> bytes:
    stream = io.BytesIO()
    canvas = Canvas(stream, pagesize=A4)
    width, height = A4

    def line(text: str, y: float, size: int = 11, bold: bool = False):
        canvas.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        canvas.drawString(55, y, text)

    line("Tahliliy hisobot", height - 60, 18, True)
    line("Murojaatlar oqimi va ko'rib chiqish holati", height - 85, 10)
    line("1. Maqsad", height - 125, 14, True)
    line("2026-yil II chorakdagi murojaatlar oqimi va ko'rib chiqish tezligini tahlil qilish.", height - 150)
    line("2. Asosiy ko'rsatkichlar", height - 190, 14, True)
    for index, value in enumerate((
        "Jami murojaatlar: 128 ta", "Ko'rib chiqilgan: 96 ta (75%)",
        "Jarayonda: 24 ta", "Qo'shimcha ma'lumot kutilmoqda: 8 ta",
        "O'rtacha ko'rib chiqish muddati: 11 kun",
    )):
        line("- " + value, height - 220 - (index * 24))
    line("3. Kuzatuvlar", height - 370, 14, True)
    line("Kelishuvlarga oid murojaatlar soni aprel oyidagi 9 tadan iyunda 17 taga oshgan.", height - 400)
    line("Eng ko'p murojaat chakana savdo (42 ta) va transport xizmatlari (31 ta) yo'nalishida.", height - 425)
    line("4. Tavsiyalar", height - 465, 14, True)
    line("- Kechikayotgan ishlar uchun haftalik nazorat ro'yxatini yuritish.", height - 495)
    line("- Takroriy mavzular bo'yicha standart ma'lumotnoma shablonini qo'llash.", height - 520)
    canvas.showPage()
    line("Tahliliy hisobot (davomi)", height - 60, 16, True)
    line("5. Oylik dinamika", height - 105, 14, True)
    rows = (("Aprel", 36, 30, 6), ("May", 43, 32, 11), ("Iyun", 49, 34, 15))
    y = height - 140
    line("Oy        Jami        Ko'rib chiqilgan        Jarayonda", y, 11, True)
    for month, total, done, progress in rows:
        y -= 28
        line(f"{month:<10} {total:<12} {done:<24} {progress}", y)
    line("Xulosa", y - 55, 14, True)
    line("Ma'lumotlar murojaatlar soni oshganini, yakunlash ulushi esa 75 foiz ekanini ko'rsatadi.", y - 85)
    canvas.save()
    return stream.getvalue()


def make_contradiction() -> bytes:
    stream = io.BytesIO()
    doc = WordDocument()
    doc.add_heading("Topshiriq bayonnomasi", 1)
    doc.add_heading("1. Muddat", 2)
    doc.add_paragraph("Topshiriqning yakuniy muddati 15-avgust 2026-yil deb belgilandi.")
    doc.add_heading("2. Mas’ullar", 2)
    doc.add_paragraph("Tahliliy bo‘lim ma’lumotlarni yig‘adi, huquqiy bo‘lim asoslarni tekshiradi.")
    doc.add_heading("3. Yakuniy kelishuv", 2)
    doc.add_paragraph("Yig‘ilish yakunida topshiriqning yakuniy muddati 20-avgust 2026-yil ekani qayd etildi.")
    doc.add_paragraph("Ikki sana o‘rtasidagi tafovut aniqlashtirilishi lozim.")
    doc.save(stream)
    return stream.getvalue()


def make_statistics() -> bytes:
    stream = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Murojaatlar"
    sheet.append(["Oy", "Murojaatlar soni", "Ko‘rib chiqilgan", "Jarayonda"])
    sheet.append(["Aprel", 36, 30, 6])
    sheet.append(["May", 43, 32, 11])
    sheet.append(["Iyun", 49, 34, 15])
    second = workbook.create_sheet("Yo‘nalishlar")
    second.append(["Yo‘nalish", "Soni"])
    second.append(["Chakana savdo", 42])
    second.append(["Transport xizmatlari", 31])
    second.append(["Davlat xaridlari", 22])
    workbook.save(stream)
    return stream.getvalue()


def generate_files() -> dict[str, bytes]:
    return {
        "murojaat_raqobat.docx": make_appeal(),
        "tahliliy_hisobot.pdf": make_report(),
        "topshiriq_bayonnomasi.docx": make_contradiction(),
        "murojaatlar_statistikasi.xlsx": make_statistics(),
    }


def ensure_user(db, username: str, full_name: str, password: str, role: Role) -> User:
    user = db.scalar(select(User).where(User.username == username))
    if user:
        # An existing account keeps its own password hash. The seed never resets
        # credentials that a real deployment or an administrator already set.
        return user
    user = User(username=username, full_name=full_name, password_hash=hash_password(password), role=role)
    db.add(user)
    db.flush()
    return user


def reset_operational_samples(db, assets: dict[str, bytes]) -> None:
    # Never delete a real user upload merely because it resembles a seeded fixture.
    # Legacy seeded rows are recognized only by their reserved stored-name prefix.
    candidates = list(db.scalars(select(Document).where(
        (Document.seed_key.like(f"{SEED_KEY_PREFIX}%")) |
        (Document.seed_key.like("raqobat-ai:v1:%")) |
        (Document.stored_name.in_([f"seed_{name}" for name in OPERATIONAL_FILES])) |
        (Document.filename.in_(OPERATIONAL_FILES))
    )))
    seeded_documents = []
    expected_hashes = {name: hashlib.sha256(data).hexdigest() for name, data in assets.items()}
    expected_texts = {
        name: re.sub(r"\s+", " ", parse_file(OPERATIONAL_FILES[name][0], data)).strip()
        for name, data in assets.items()
    }
    for document in candidates:
        path = get_settings().data_dir / "documents" / document.stored_name
        exact_fixture_copy = (
            document.filename in expected_hashes and path.exists()
            and (
                hashlib.sha256(path.read_bytes()).hexdigest() == expected_hashes[document.filename]
                or re.sub(r"\s+", " ", document.parsed_text).strip() == expected_texts[document.filename]
            )
        )
        if document.seed_key or document.stored_name.startswith("seed_") or exact_fixture_copy:
            seeded_documents.append(document)
    for document in seeded_documents:
        path = get_settings().data_dir / "documents" / document.stored_name
        db.delete(document)
        if path.exists():
            path.unlink()
    db.execute(delete(Task).where(Task.seed_key.like(f"{SEED_KEY_PREFIX}%")))
    db.flush()


def seed(args) -> None:
    if not args.confirm_development:
        raise SystemExit("Operatsion namuna faqat --confirm-development bayrog‘i bilan ishlaydi")
    assets = generate_files()
    sample_dir = PROJECT_ROOT / "sample-data"
    sample_dir.mkdir(exist_ok=True)
    for filename, data in assets.items():
        (sample_dir / filename).write_bytes(data)

    settings = get_settings()
    target = settings.data_dir / "documents"
    target.mkdir(parents=True, exist_ok=True)
    with SessionLocal() as db:
        if args.reset:
            reset_operational_samples(db, assets)
            db.commit()
            print("Oldingi operatsion namuna yozuvlari tozalandi")
            if args.reset_only:
                return
        admin = db.scalar(select(User).where(User.username == args.admin_username))
        if not admin or admin.role != Role.administrator:
            raise SystemExit(f"Administrator topilmadi: {args.admin_username}")
        # Deliberately simple credentials are allowed only by this explicitly
        # confirmed local/development seed command, and only for accounts it
        # creates itself. Production users are never created by application
        # startup and must use managed secrets.
        seed_password = getattr(args, "seed_password", "") or get_settings().local_seed_password
        if len(seed_password) < 8:
            raise SystemExit("Lokal seed paroli kamida 8 belgidan iborat bo‘lishi kerak")
        rahbar = ensure_user(db, "rahbar_analitika", "Analitika bo‘limi rahbari", seed_password, Role.rahbar)
        xodim = ensure_user(db, "xodim_huquq", "Huquqiy tahlil bo‘yicha xodim", seed_password, Role.xodim)
        db.flush()

        created_documents = 0
        for filename, (kind, category) in OPERATIONAL_FILES.items():
            seed_key = f"{SEED_KEY_PREFIX}{filename}"
            exists = db.scalar(select(Document).where(Document.seed_key == seed_key))
            if exists:
                continue
            data = assets[filename]
            parsed = parse_file(kind, data)
            stored_name = f"seed_{filename}"
            (target / stored_name).write_bytes(data)
            document = Document(
                owner_id=xodim.id, filename=filename, stored_name=stored_name,
                media_type=MIMES[kind], size_bytes=len(data), parsed_text=parsed,
                category=category, is_confidential=False,
                seed_key=seed_key,
            )
            db.add(document)
            db.flush()
            index_document(db, document, "document")
            created_documents += 1

        now = datetime.now(timezone.utc)
        created_tasks = 0
        for index, (title, task_status, days) in enumerate(TASKS):
            if db.scalar(select(Task).where(Task.title == title)):
                continue
            assignee = xodim if index % 3 else rahbar
            db.add(Task(
                title=title,
                description="Ichki ish jarayonida nazorat qilinadigan topshiriq.",
                assigned_to=assignee.id, created_by=admin.id, status=task_status,
                deadline=now + timedelta(days=days), seed_key=f"{SEED_KEY_PREFIX}task:{index}",
            ))
            created_tasks += 1
        db.commit()
        print(f"Operatsion namuna tayyor: {created_documents} yangi hujjat, {created_tasks} yangi topshiriq")


def main():
    parser = argparse.ArgumentParser(description="Raqobat AI uchun idempotent operatsion namuna seed")
    parser.add_argument("--confirm-development", action="store_true")
    parser.add_argument("--admin-username", default="admin")
    parser.add_argument("--reset", action="store_true")
    parser.add_argument("--reset-only", action="store_true")
    parser.add_argument("--seed-password", default="",
                        help="Faqat yangi yaratiladigan lokal seed hisoblari uchun parol "
                             "(standart: LOCAL_SEED_PASSWORD)")
    seed(parser.parse_args())


if __name__ == "__main__":
    main()
