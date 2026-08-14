import io
import re
import uuid
import zipfile
from pathlib import Path

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from openpyxl import load_workbook
from pypdf import PdfReader

from ..config import get_settings

MIMES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


def detect_type(data: bytes) -> str:
    if data.startswith(b"%PDF-"):
        return "pdf"
    if data.startswith(b"PK"):
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                names = set(archive.namelist())
                if "word/document.xml" in names:
                    return "docx"
                if "xl/workbook.xml" in names:
                    return "xlsx"
        except zipfile.BadZipFile:
            pass
    raise HTTPException(status.HTTP_415_UNSUPPORTED_MEDIA_TYPE, "Fayl turi qo‘llab-quvvatlanmaydi yoki fayl buzilgan")


def parse_file(kind: str, data: bytes) -> str:
    try:
        if kind == "pdf":
            reader = PdfReader(io.BytesIO(data))
            parts = [f"[Sahifa {i}]\n{page.extract_text() or ''}" for i, page in enumerate(reader.pages, 1)]
            text = "\n\n".join(parts)
        elif kind == "docx":
            doc = DocxDocument(io.BytesIO(data))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells]
                    if any(values):
                        parts.append(" | ".join(values))
            text = "\n\n".join(parts)
        else:
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"[Varaq: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(v).strip() if v is not None else "" for v in row]
                    if any(values):
                        parts.append(" | ".join(values))
            text = "\n".join(parts)
    except Exception as exc:
        raise HTTPException(status.HTTP_422_UNPROCESSABLE_CONTENT, "Fayl matnini o‘qib bo‘lmadi") from exc
    text = re.sub(r"[ \t]+", " ", text).strip()
    if not text:
        raise HTTPException(status.HTTP_422_UNPROCESSABLE_CONTENT, "Faylda o‘qiladigan matn topilmadi")
    return text


async def receive_upload(file: UploadFile, folder: str) -> tuple[str, str, bytes, str]:
    settings = get_settings()
    data = await file.read(settings.max_upload_mb * 1024 * 1024 + 1)
    if len(data) > settings.max_upload_mb * 1024 * 1024:
        raise HTTPException(status.HTTP_413_CONTENT_TOO_LARGE, f"Fayl {settings.max_upload_mb} MB dan katta bo‘lmasligi kerak")
    kind = detect_type(data)
    ext = Path(file.filename or "").suffix.lower().lstrip(".")
    if ext and ext != kind:
        raise HTTPException(status.HTTP_415_UNSUPPORTED_MEDIA_TYPE, "Fayl kengaytmasi uning haqiqiy turiga mos emas")
    parsed = parse_file(kind, data)
    target_dir = settings.data_dir / folder
    target_dir.mkdir(parents=True, exist_ok=True)
    stored_name = f"{uuid.uuid4()}.{kind}"
    (target_dir / stored_name).write_bytes(data)
    return kind, stored_name, data, parsed
