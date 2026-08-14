import io
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..config import get_settings

GREEN = RGBColor(18, 102, 78)
MUTED = RGBColor(91, 105, 101)
# Outgoing correspondence follows institutional stationery conventions rather than
# the in-app brand palette: near-black body ink, restrained grey for secondary lines.
INK = RGBColor(26, 26, 26)
INK_HEX = "1A1A1A"
# A4 minus the official-letter side margins configured in `_configure`.
LETTER_TEXT_WIDTH = Cm(21) - Cm(2.5) - Cm(1.8)
# Adresat block sits in the upper-right quadrant, left-aligned inside its own column.
RECIPIENT_INDENT = Cm(9.5)


def _font(run, size: float = 11, bold: bool | None = None, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "uz-Latn-UZ")
    run._element.get_or_add_rPr().append(lang)


def _page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    _font(run, 9, False, MUTED)


def _configure(doc: Document, *, official_letter: bool = False, profile=None) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8 if official_letter else 2.2)
    section.bottom_margin = Cm(1.8 if official_letter else 2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after in (
        ("Title", 22, 0, 14), ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6), ("Heading 3", 11.5, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = GREEN
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer_text = _profile_value(profile, "footer_text") if official_letter else ""
    if footer_text:
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_after = Pt(1)
        _font(footer.add_run(footer_text), 8, False, MUTED)
        page_footer = section.footer.add_paragraph()
        page_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        page_footer.paragraph_format.space_before = Pt(0)
        _page_number(page_footer)
    else:
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _page_number(footer)

    settings = doc.settings._element
    for node in settings.findall(qn("w:compat")):
        settings.remove(node)
    compat = OxmlElement("w:compat")
    item = OxmlElement("w:compatSetting")
    item.set(qn("w:name"), "compatibilityMode")
    item.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    item.set(qn("w:val"), "15")
    compat.append(item)
    settings.append(compat)


def _add_rich_text(paragraph, value: str) -> None:
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", value):
        if match.start() > cursor:
            _font(paragraph.add_run(value[cursor:match.start()]))
        _font(paragraph.add_run(match.group(1)), bold=True)
        cursor = match.end()
    if cursor < len(value):
        _font(paragraph.add_run(value[cursor:]))


def _add_hyperlink(paragraph, text: str, url: str, color_hex: str = "12664E") -> None:
    relation = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), color_hex)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    props.extend([color, underline]); run.append(props)
    node = OxmlElement("w:t"); node.text = text; run.append(node); hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_markdown(doc: Document, content: str) -> None:
    for raw in content.splitlines():
        line = raw.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            doc.add_heading(re.sub(r"\*\*", "", heading.group(2)), level=len(heading.group(1)))
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        numbered = re.match(r"^\d+[.)]\s+(.+)$", line)
        if bullet or numbered:
            paragraph = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            paragraph.paragraph_format.left_indent = Cm(0.65)
            paragraph.paragraph_format.first_line_indent = Cm(-0.35)
            paragraph.paragraph_format.space_after = Pt(4)
            _add_rich_text(paragraph, (bullet or numbered).group(1))
            continue
        paragraph = doc.add_paragraph()
        if line.startswith(">"):
            line = line.lstrip("> ")
            paragraph.paragraph_format.left_indent = Cm(0.65)
            paragraph.paragraph_format.right_indent = Cm(0.4)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
        _add_rich_text(paragraph, line)


def _profile_value(profile, name: str) -> str:
    return str(getattr(profile, name, "") or "").strip()


def _official_date(value) -> str:
    """Render a correspondence date the way Uzbek official stationery does (kk.oo.yyyy).

    The stored value stays ISO; only the printed letter is localized. An unset date
    keeps the explicit `[sana]` placeholder so a draft can never look finalized.
    """
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    text = str(value or "").strip()
    if not text:
        return "[sana]"
    iso = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", text)
    return f"{iso.group(3)}.{iso.group(2)}.{iso.group(1)}" if iso else text


# Internal evidence identifiers handed to the generator (D1 = document evidence,
# L1 = legal evidence). They index the internal evidence sheet and must never
# appear in the recipient-facing letter.
_CITATION_TOKEN_RE = re.compile(r"\s*\[(?:\d+(?:\s*,\s*\d+)*)\]")
_BRACKETED_EVIDENCE_ID_RE = re.compile(
    r"\s*[(\[]\s*[DL]\d{1,2}(?:\s*[,;]\s*[DL]\d{1,2})*\s*[)\]]"
)
_BARE_EVIDENCE_ID_RE = re.compile(r"(?<![\w-])[DL]\d{1,2}(?![\w-])")


def _external_text(value: str) -> str:
    """Remove internal citation tokens and draft-only placeholder residue."""
    text = _CITATION_TOKEN_RE.sub("", value or "")
    text = _BRACKETED_EVIDENCE_ID_RE.sub("", text)
    text = _BARE_EVIDENCE_ID_RE.sub("", text)
    # Stripping a mid-sentence token leaves doubled spaces or a space before
    # punctuation; the letter must still read as ordinary formal prose.
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    return text.strip()


def _law_title(source: dict) -> str:
    title = re.sub(r"\s*\([^)]*(?:O['‘’]?RQ|ЎРҚ|№)[^)]*\)\s*$", "", source.get("document_name", ""),
                   flags=re.IGNORECASE).strip()
    title = re.sub(r"\s+(?:to['‘’]?g['‘’]?risida(?:gi)?\s+)?qonuni?$", "", title, flags=re.IGNORECASE).strip()
    return title or source.get("document_name", "Normativ-huquqiy hujjat")


def _formal_legal_reference(source: dict) -> str:
    article = source.get("display_label") or source.get("article_or_clause") or "tegishli normasi"
    article = re.sub(r"\s+", " ", article).strip()
    article = re.sub(r"^(\d+)\s*[-.]?\s*modda$", r"\1-moddasi", article, flags=re.IGNORECASE)
    article = re.sub(r"^(\d+)\s*[-.]?\s*band$", r"\1-bandi", article, flags=re.IGNORECASE)
    category = (source.get("document_type") or "").lower()
    if category == "qonun":
        return f"O‘zbekiston Respublikasining “{_law_title(source)}”gi Qonunining {article}"
    return f"“{source.get('document_name', 'Normativ-huquqiy hujjat')}”ning {article}"


def _add_letterhead(doc: Document, profile) -> None:
    logo_name = _profile_value(profile, "logo_stored_name")
    if logo_name:
        logo_path = get_settings().data_dir / "organization" / Path(logo_name).name
        if logo_path.exists():
            try:
                doc.add_picture(str(logo_path), height=Cm(1.55))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.paragraphs[-1].paragraph_format.space_after = Pt(3)
            except Exception:
                # A broken optional logo must never block a legally grounded export.
                pass
    parent = _profile_value(profile, "parent_organization")
    if parent:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        _font(p.add_run(parent.upper()), 8.5, True, MUTED)
    name = _profile_value(profile, "organization_name")
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        _font(p.add_run(name.upper()), 13, True, INK)
    secondary_name = _profile_value(profile, "organization_name_secondary")
    if secondary_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        _font(p.add_run(secondary_name.upper()), 8.5, True, MUTED)
    department = _profile_value(profile, "department")
    if department:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        _font(p.add_run(department), 9.5, True, MUTED)
    letterhead_text = _profile_value(profile, "letterhead_text")
    if letterhead_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        _font(p.add_run(letterhead_text), 8.5, False, MUTED)
    contacts = [value for value in (
        _profile_value(profile, "address"), _profile_value(profile, "phone"),
        _profile_value(profile, "email"), _profile_value(profile, "website"),
        f"STIR: {_profile_value(profile, 'tax_id')}" if _profile_value(profile, "tax_id") else "",
    ) if value]
    if contacts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        _font(p.add_run("  •  ".join(contacts)), 8.5, False, MUTED)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), INK_HEX)
    borders.append(bottom); rule._p.get_or_add_pPr().append(borders)


def _add_structured_letter(doc: Document, value: dict, profile, sources: list[dict], *, draft: bool) -> None:
    _add_letterhead(doc, profile)
    if draft:
        marker = doc.add_paragraph()
        marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        marker.paragraph_format.space_after = Pt(6)
        required = ("organization_name", "address", "phone", "email", "website",
                    "outgoing_prefix", "signatory_name", "signatory_title")
        missing_profile = any(not _profile_value(profile, field) for field in required)
        marker_text = (
            "QORALAMA — RASMIY REKVIZITLAR TO‘LDIRILMAGAN; YUBORISH MUMKIN EMAS"
            if missing_profile else
            "QORALAMA — YUBORISHDAN OLDIN VAKOLATLI XODIM TEKSHIRUVI TALAB ETILADI"
        )
        _font(marker.add_run(marker_text),
              8.5, True, MUTED)
    reference = doc.add_paragraph()
    reference.alignment = WD_ALIGN_PARAGRAPH.LEFT
    prefix = _profile_value(profile, "outgoing_prefix")
    document_date = _official_date(value.get("document_date"))
    outgoing = str(value.get("outgoing_number") or "").strip() or "[chiqish raqami]"
    number = f"{prefix}-{outgoing}" if prefix and outgoing != "[chiqish raqami]" and not outgoing.startswith(prefix) else outgoing
    _font(reference.add_run(f"{document_date}    №  {number}"), 10.5)
    # Adresat block: left-aligned inside an indented right-hand column so a
    # multi-line recipient wraps as a block instead of ragged right-aligned lines.
    recipient = doc.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.LEFT
    recipient.paragraph_format.left_indent = RECIPIENT_INDENT
    recipient.paragraph_format.space_before = Pt(6)
    _font(recipient.add_run(str(value.get("recipient") or "").strip()
                            or "[Qabul qiluvchi tashkilot yoki shaxs]"), 11, True)
    subject = doc.add_paragraph()
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject.paragraph_format.space_before = Pt(8)
    subject.paragraph_format.space_after = Pt(9)
    _font(subject.add_run("Mavzu: "), 11, True)
    _font(subject.add_run(value.get("subject", "")))
    salutation = doc.add_paragraph()
    salutation.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(salutation.add_run(value.get("salutation", "")), bold=True)
    for item in value.get("appeal_summary", []):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25); p.paragraph_format.space_after = Pt(5)
        _font(p.add_run(_external_text(item)))
    legal_sources = [source for source in sources if source.get("evidence_type") == "nhh"]
    for item in value.get("legal_basis", []):
        ids = item.get("source_ids") or []
        selected = []
        for source_id in ids:
            match = re.fullmatch(r"L(\d+)", str(source_id))
            if match and int(match.group(1)) <= len(legal_sources):
                selected.append(legal_sources[int(match.group(1)) - 1])
        statement = _external_text(item.get("statement", ""))
        references = "; ".join(_formal_legal_reference(source) for source in selected)
        if references:
            if statement.lower().startswith(("tegishli norma", "ushbu norma", "mazkur norma")):
                statement = re.sub(r"^(?:tegishli|ushbu|mazkur)\s+norma", references, statement,
                                   count=1, flags=re.IGNORECASE)
            elif references.lower() not in statement.lower():
                statement = f"{references}ga muvofiq, {statement[:1].lower() + statement[1:]}"
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25); p.paragraph_format.space_after = Pt(5)
        _font(p.add_run(statement))
    for item in value.get("conclusion", []):
        conclusion = _external_text(item)
        if legal_sources:
            conclusion = re.sub(
                r"\b\d{1,3}\s*[-.]?\s*modda(?:si)?\b",
                _formal_legal_reference(legal_sources[0]),
                conclusion,
                count=1,
                flags=re.IGNORECASE,
            )
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25); p.paragraph_format.space_after = Pt(5)
        _font(p.add_run(conclusion))
    closing = _external_text(value.get("closing", ""))
    if closing and "[" not in closing:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(7); _font(p.add_run(closing))
    spacer = doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(2)
    # Signature block: post on the left, signer flush to the right margin via an
    # explicit right tab stop (default tab stops leave the name floating mid-line).
    sign = doc.add_paragraph()
    sign.paragraph_format.space_before = Pt(10)
    sign.paragraph_format.tab_stops.add_tab_stop(LETTER_TEXT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    title = _profile_value(profile, "signatory_title") or "[Lavozim]"
    person = _profile_value(profile, "signatory_name") or "[F.I.Sh.]"
    _font(sign.add_run(title), 11, True)
    sign.add_run("\t")
    _font(sign.add_run(person), 11, True)
    verification = _profile_value(profile, "qr_verification_url")
    barcode = _profile_value(profile, "barcode_text")
    if verification or barcode:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
        if verification:
            _font(p.add_run("Elektron tekshiruv: "), 8, True, MUTED)
            _add_hyperlink(p, verification, verification, INK_HEX)
        if barcode:
            _font(p.add_run(("    " if verification else "") + f"Identifikator: {barcode}"), 8, False, MUTED)


def _add_sources(doc: Document, sources: list[dict]) -> None:
    if not sources:
        return
    groups = (
        ("document", "Asos bo‘lgan hujjat dalillari"),
        ("nhh", "Foydalanilgan normativ-huquqiy manbalar"),
    )
    for evidence_type, heading in groups:
        values = [source for source in sources if source.get("evidence_type", "document") == evidence_type]
        if not values:
            continue
        doc.add_heading(heading, level=1)
        for source in values:
            p = doc.add_paragraph()
            _font(p.add_run(f"[{source.get('citation_number')}] {source['document_name']}"), bold=True)
            location = []
            if evidence_type == "nhh" and source.get("display_label"):
                location.append(source["display_label"])
            if evidence_type == "document" and source.get("section"):
                location.append(source["section"])
            if source.get("page"):
                location.append(f"{source['page']}-sahifa")
            if location:
                loc = doc.add_paragraph(); _font(loc.add_run("Joylashuv: "), bold=True)
                _font(loc.add_run(" • ".join(location)), color=MUTED)
            url = source.get("url") or ""
            if url.startswith(("http://", "https://")):
                link = doc.add_paragraph(); _add_hyperlink(link, "Rasmiy manbani ochish", url)


def make_docx(title: str, content: str, sources: list[dict],
              structured: dict | None = None, profile=None, *, draft: bool = True) -> bytes:
    doc = Document()
    official_letter = bool(structured and structured.get("kind") == "response_letter")
    _configure(doc, official_letter=official_letter, profile=profile)
    if official_letter:
        _add_structured_letter(doc, structured, profile, sources, draft=draft)
    else:
        title_paragraph = doc.add_paragraph(style="Title")
        _font(title_paragraph.add_run(title), 22, True, GREEN)
        _add_markdown(doc, content)
        _add_sources(doc, sources)
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def make_internal_evidence_docx(title: str, sources: list[dict]) -> bytes:
    doc = Document()
    _configure(doc)
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(4)
    _font(heading.add_run("ICHKI DALILLAR VARAQASI"), 16, True, GREEN)
    note = doc.add_paragraph()
    _font(note.add_run(
        "Xizmat doirasida foydalanish uchun. Ushbu hujjat recipientga yuboriladigan xat tarkibiga kirmaydi."
    ), 9.5, False, MUTED)
    subject = doc.add_paragraph(); _font(subject.add_run(title), 11, True)
    _add_sources(doc, sources)
    stream = io.BytesIO(); doc.save(stream)
    return stream.getvalue()
